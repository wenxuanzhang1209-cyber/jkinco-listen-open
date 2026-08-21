"""报告导出:DOCX 与 PDF。

从 JKincoListen.py 单体抽出。优先套用 templates/v2 下的定制版式模板生成 DOCX,
模板不可用时回退到内置排版;PDF 走 reportlab 直排并做中文换行。
导出文件写入临时目录,由 backend 在下载完成后删除(见 backend/main.py 的导出端点)。

注:APP_ROOT 在此独立定义。本文件与 JKincoListen.py 同在项目根目录,
Path(__file__).resolve().parent 取值与单体中的定义完全一致,避免循环导入。
"""
from __future__ import annotations

import os
import re
import tempfile
import uuid
import time
from pathlib import Path

from jkinco_scenes import (
    is_auto_mode,
    is_customer_visit_mode,
    is_interview_mode,
    is_personal_mode,
    is_talk_mode,
    output_title,
)
from jkinco_text import clean_markdown_text, markdown_lines, wrap_pdf_line

from jkinco_logging import get_logger

LOGGER = get_logger("export")

APP_ROOT = Path(__file__).resolve().parent


TALK_TEMPLATE_PATH = Path(os.getenv(
    "JKINCO_TALK_TEMPLATE",
    str(APP_ROOT / "templates" / "v2" / "会议纪要_工地例会原版式模板.docx"),
))


REPORT_LOGO_PATH = Path(os.getenv(
    "JKINCO_REPORT_LOGO",
    str(APP_ROOT / "assets" / "sribs-meeting-logo.jpeg"),
))


PERSONAL_TEMPLATE_PATH = Path(os.getenv(
    "JKINCO_PERSONAL_TEMPLATE",
    str(APP_ROOT / "templates" / "个人备忘录总结_会议纪要模板.docx"),
))


INTERVIEW_TEMPLATE_PATH = Path(os.getenv(
    "JKINCO_INTERVIEW_TEMPLATE",
    str(APP_ROOT / "templates" / "HR面试记录与候选人反馈模板.docx"),
))


CUSTOMER_VISIT_TEMPLATE_PATH = Path(os.getenv(
    "JKINCO_CUSTOMER_VISIT_TEMPLATE",
    str(APP_ROOT / "templates" / "客户拜访会议纪要模板_建科.docx"),
))


EXPORT_DIR = Path(tempfile.gettempdir()) / "jkinco_exports"


def export_filename(suffix):
    """生成导出文件的落盘路径。

    名字里必须带随机串:原先只用「秒 + 毫秒」,而生成文件名本身只花微秒 ——
    实测 200 次并发只得到 5 个不同名字,碰撞 195 次。导出目录是全体用户共用的,
    撞名意味着后写的覆盖先写的,先导出的那位下载到的是别人的会议纪要。
    """
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    unique = uuid.uuid4().hex[:12]
    return str(EXPORT_DIR / f"筑听结构化纪要_{timestamp}_{unique}.{suffix}")


def export_summary_docx(summary_text, app_mode="auto"):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    try:
        from report_templates import build_report_docx

        path = Path(export_filename("docx"))
        build_report_docx(
            clean_markdown_text(summary_text),
            app_mode,
            path,
            talk_template=TALK_TEMPLATE_PATH,
            logo=REPORT_LOGO_PATH,
        )
        return str(path)
    except Exception as error:
        LOGGER.warning("定制 DOCX 模板生成失败,降级为兼容导出:%s", error)

    def set_cn_font(run, font_name="Hiragino Sans GB"):
        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), font_name)

    def template_path_for_mode(mode):
        if is_talk_mode(mode):
            return TALK_TEMPLATE_PATH
        if is_personal_mode(mode):
            return PERSONAL_TEMPLATE_PATH
        if is_interview_mode(mode):
            return INTERVIEW_TEMPLATE_PATH
        if is_customer_visit_mode(mode):
            return CUSTOMER_VISIT_TEMPLATE_PATH
        return None

    def remove_report_marker(doc):
        for para in list(doc.paragraphs):
            if "{{REPORT_BODY}}" in para.text:
                para._element.getparent().remove(para._element)

    def add_text_paragraph(doc, text):
        bullet_match = re.match(r"^[-*•]\s+(.*)$", text)
        numbered_match = re.match(r"^\d+[.、]\s*(.*)$", text)
        if bullet_match:
            para = doc.add_paragraph(bullet_match.group(1), style="List Bullet")
        elif numbered_match:
            para = doc.add_paragraph(numbered_match.group(1), style="List Number")
        else:
            para = doc.add_paragraph(text)
        for run in para.runs:
            set_cn_font(run)
        return para

    def add_pipe_table(doc, rows):
        cells = [
            [cell.strip() or " " for cell in row.strip().strip("|").split("|")]
            for row in rows
        ]
        max_cols = max((len(row) for row in cells), default=1)
        table = doc.add_table(rows=len(cells), cols=max_cols)
        table.style = "Table Grid"
        for row_index, row in enumerate(cells):
            for col_index in range(max_cols):
                value = row[col_index] if col_index < len(row) else " "
                cell = table.cell(row_index, col_index)
                cell.text = value
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_cn_font(run)
                        if row_index == 0:
                            run.bold = True
        doc.add_paragraph("")

    specialized_template = template_path_for_mode(app_mode)
    if specialized_template and specialized_template.exists():
        doc = Document(str(specialized_template))
        remove_report_marker(doc)

        generated = doc.add_paragraph()
        generated.alignment = 2
        sub_run = generated.add_run(f"AI 自动整理｜生成日期：{time.strftime('%Y年%m月%d日')}")
        sub_run.font.size = Pt(8.5)
        sub_run.font.color.rgb = RGBColor(100, 116, 139)
        set_cn_font(sub_run)

        table_buffer = []
        for raw_line in markdown_lines(summary_text):
            line = raw_line.strip()
            if line.replace(" ", "") == output_title(app_mode).replace(" ", ""):
                continue
            if not line:
                if table_buffer:
                    add_pipe_table(doc, table_buffer)
                    table_buffer = []
                doc.add_paragraph("")
                continue
            if "|" in line and not line.startswith("#"):
                table_buffer.append(line)
                continue
            if table_buffer:
                add_pipe_table(doc, table_buffer)
                table_buffer = []
            if line.startswith("#"):
                level = min(len(line) - len(line.lstrip("#")), 3)
                text = line.lstrip("#").strip()
                para = doc.add_heading(text, level=level)
                for run in para.runs:
                    set_cn_font(run)
                continue
            chinese_heading = re.match(r"^[一二三四五六七八九十]+、", line)
            if chinese_heading:
                para = doc.add_heading(line, level=1)
                for run in para.runs:
                    set_cn_font(run)
                continue
            add_text_paragraph(doc, line)
        if table_buffer:
            add_pipe_table(doc, table_buffer)

        path = export_filename("docx")
        doc.save(path)
        return path

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("筑听结构化纪要")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 99, 182)
    set_cn_font(run)

    meta = doc.add_paragraph(f"导出时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    set_cn_font(meta.runs[0])

    for raw_line in markdown_lines(summary_text):
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            text = line.lstrip("#").strip()
            para = doc.add_heading(text, level=level)
            for run in para.runs:
                set_cn_font(run)
            continue
        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        numbered_match = re.match(r"^\d+[.、]\s*(.*)$", line)
        if bullet_match:
            para = doc.add_paragraph(bullet_match.group(1), style="List Bullet")
        elif numbered_match:
            para = doc.add_paragraph(numbered_match.group(1), style="List Number")
        else:
            para = doc.add_paragraph(line)
        for run in para.runs:
            set_cn_font(run)

    path = export_filename("docx")
    doc.save(path)
    return path


def export_summary_pdf(summary_text, app_mode="auto"):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    try:
        from report_templates import build_report_docx, convert_docx_to_pdf

        docx_path = Path(export_filename("docx"))
        pdf_path = Path(export_filename("pdf"))
        build_report_docx(
            clean_markdown_text(summary_text),
            app_mode,
            docx_path,
            talk_template=TALK_TEMPLATE_PATH,
            logo=REPORT_LOGO_PATH,
        )
        if convert_docx_to_pdf(docx_path, pdf_path):
            docx_path.unlink(missing_ok=True)
            return str(pdf_path)
        docx_path.unlink(missing_ok=True)
        LOGGER.warning("未检测到可用的 LibreOffice,降级为兼容 PDF 导出")
    except Exception as error:
        LOGGER.warning("定制 PDF 模板生成失败,降级为兼容导出:%s", error)

    font_name = "STSong-Light"
    pdfmetrics.registerFont(UnicodeCIDFont(font_name))

    path = export_filename("pdf")
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    left = 22 * mm
    right = 22 * mm
    top = height - 24 * mm
    y = top
    max_width = width - left - right

    def new_page_if_needed(line_height=16):
        nonlocal y
        if y < 24 * mm + line_height:
            c.showPage()
            c.setFont(font_name, 11)
            y = top

    c.setFont(font_name, 18)
    c.setFillColorRGB(0.04, 0.34, 0.65)
    c.drawString(left, y, output_title(app_mode) if not is_auto_mode(app_mode) else "筑听结构化纪要")
    y -= 18
    c.setFont(font_name, 9)
    c.setFillColorRGB(0.39, 0.45, 0.54)
    c.drawString(left, y, f"导出时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
    y -= 22

    c.setFillColorRGB(0.07, 0.12, 0.20)
    for raw_line in markdown_lines(summary_text):
        line = raw_line.strip()
        if not line:
            y -= 8
            continue
        heading_level = len(line) - len(line.lstrip("#")) if line.startswith("#") else 0
        if heading_level:
            text = line.lstrip("#").strip()
            font_size = 14 if heading_level == 1 else 12
            c.setFont(font_name, font_size)
            c.setFillColorRGB(0.04, 0.34, 0.65)
            y -= 4
            for wrapped in wrap_pdf_line(text, c, font_name, font_size, max_width):
                new_page_if_needed(font_size + 6)
                c.drawString(left, y, wrapped)
                y -= font_size + 6
            c.setFillColorRGB(0.07, 0.12, 0.20)
            continue
        line = re.sub(r"^[-*]\s+", "• ", line)
        c.setFont(font_name, 10.5)
        for wrapped in wrap_pdf_line(line, c, font_name, 10.5, max_width):
            new_page_if_needed(15)
            c.drawString(left, y, wrapped)
            y -= 15
    c.save()
    return path
