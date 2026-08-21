"""Secure, user-owned DOCX templates for minutes generation and export.

The source DOCX is immutable. Every export starts from the stored source bytes,
which makes regeneration idempotent and prevents content from accumulating
across repeated exports.
"""
from __future__ import annotations

import hashlib
import io
import os
import json
import re
import sqlite3
import time
import uuid
import zipfile
from pathlib import Path, PurePosixPath
from typing import Any, Iterable
from xml.etree import ElementTree

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from backend.auth import PROFILE_DB, PROFILE_DB_LOCK, is_guest
from jkinco_llm import call_llm
from jkinco_prompts import SOURCE_GUARD, fence_source
from jkinco_text import BIDI_OVERRIDES, clean_markdown_text, strip_control_characters, user_facing_error

MAX_TEMPLATE_BYTES = 10 * 1024 * 1024
# 每个账号的模板配额。模板正文是直接存进 SQLite 的 BLOB,单个上限 10MB,原先
# 数量与总量都不设限 —— 一个账号(含免注册访客)反复上传就能把磁盘写满,而磁盘
# 写满时整个平台一起停摆,不只是模板功能。数量按未删除的算,方便用户删掉腾位;
# 总量把软删除的也算进去,否则「传满—删掉—再传」的循环在 30 天保留期内无上限。
MAX_TEMPLATES_PER_USER = max(1, int(os.getenv("JKINCO_MAX_TEMPLATES_PER_USER", "20")))
MAX_TEMPLATE_STORAGE_PER_USER = max(
    MAX_TEMPLATE_BYTES, int(os.getenv("JKINCO_MAX_TEMPLATE_STORAGE_PER_USER", str(40 * 1024 * 1024)))
)
# 访客的配额单独收紧,口径与作业槽位的 MAX_PENDING_JOBS_PER_GUEST 一致。
# 按账号给配额对访客形同虚设:访客免注册,单 IP 在保留期(8 小时)内最多能开出
# 两千多个账号,每个都按 40MB 计,放大出来的总量足以写满磁盘 —— 而磁盘写满时
# 整个平台一起停摆,不只是模板功能。访客是临时通道,一两个模板已经够用。
MAX_TEMPLATES_PER_GUEST = max(1, int(os.getenv("JKINCO_MAX_TEMPLATES_PER_GUEST", "3")))
MAX_TEMPLATE_STORAGE_PER_GUEST = max(
    MAX_TEMPLATE_BYTES, int(os.getenv("JKINCO_MAX_TEMPLATE_STORAGE_PER_GUEST", str(10 * 1024 * 1024)))
)
MAX_ARCHIVE_ENTRIES = 500
MAX_UNCOMPRESSED_BYTES = 60 * 1024 * 1024
MAX_COMPRESSION_RATIO = 150
# 解压后 XML 的总量。这道限制是在解析之前生效的,因此它决定了单次上传能让
# 服务器付出的最坏 CPU/内存代价 —— 12MB 时实测最坏 21.6 万段落、11.6 秒、
# 峰值近 1GB。会议纪要模板本质是一张待填的表,3MB 已相当于 5.4 万个段落,
# 比任何真实模板宽两个数量级。
#
# 这道限制决定了最坏 CPU 代价,而原先注释里写的「3 秒」低估了一个数量级:实测
# analyze_docx 约 0.9ms/段,2 万段(document.xml 仅 1.29MB,远在限内)就要 13.5 秒,
# 按 3MB 上限外推约半分钟。所以光有这道上限不够 —— 上传接口另有按账号的频次限制,
# 且配额预检排在解析之前(见 _precheck_quota)。
MAX_XML_BYTES = max(256 * 1024, int(os.getenv("JKINCO_TEMPLATE_MAX_XML_BYTES", str(3 * 1024 * 1024))))
MAX_TEMPLATE_CONTEXT_CHARS = 12_000
MAX_PREVIEW_ITEMS = 240
# 解析结果的条数上限。归档层允许 12MB 的 document.xml,而每个段落至少产出一条
# 插入候选项 —— 实测一个 4MB 的合法 docx(20.7 万段落)会解析出 20.7 万条候选,
# 占用近 1GB 内存、落库 40MB。这些低置信度节点只是给人工下拉框用的,留几百条
# 就够;高置信度候选(占位符、语义标题)单独计数,不会被它们挤掉。
MAX_INSERTION_CANDIDATES = 240
MAX_PLACEHOLDERS = 500
REQUIRED_DOCX_ENTRIES = {"[Content_Types].xml", "word/document.xml"}
ALLOWED_TEMPLATE_MIME_TYPES = {
    "",
    "application/octet-stream",
    "application/zip",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
DANGEROUS_PART_PREFIXES = (
    "word/activex/",
    "word/embeddings/",
    "word/vbaproject",
)
DANGEROUS_RELATIONSHIP_SUFFIXES = (
    "/attachedtemplate",
    "/oleobject",
    "/package",
)
SCENARIOS = {"auto", "talk", "general", "personal", "interview", "customer_visit"}

FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    "title": ("会议主题", "会议名称", "标题", "title", "meetingtitle"),
    "meeting_type": ("会议类型", "场景", "meetingtype"),
    "date": ("会议日期", "日期", "date"),
    "start_time": ("开始时间", "会议开始时间", "starttime"),
    "end_time": ("结束时间", "会议结束时间", "endtime"),
    "location": ("会议地点", "地点", "location"),
    "participants": ("参会人员", "与会人员", "参与人", "participants", "attendees"),
    "host": ("主持人", "host"),
    "recorder": ("记录人", "纪要人", "recorder"),
    "client": ("客户名称", "客户", "client"),
    "project": ("项目名称", "项目", "project"),
    "summary": ("会议摘要", "核心摘要", "摘要", "summary"),
    "minutes": ("会议纪要", "纪要正文", "完整纪要", "正文", "minutes", "content"),
    "agenda": ("会议议题", "议题", "agenda"),
    "conclusions": ("会议结论", "结论", "conclusions"),
    "todos": ("待办事项", "行动项", "任务清单", "todos", "actions"),
    "transcript": ("原始转写", "转写文本", "transcript"),
    "generated_at": ("生成时间", "导出时间", "generatedat"),
}
LIST_FIELDS = {"agenda", "conclusions", "todos"}
STRUCTURED_FIELD_LABELS = {
    "task": "待办事项",
    "todo": "待办事项",
    "action": "行动项",
    "description": "事项说明",
    "topic": "议题",
    "agenda": "议题",
    "conclusion": "结论",
    "owner": "负责人",
    "responsible": "负责人",
    "responsible_party": "责任方",
    "deadline": "截止时间",
    "due_date": "截止时间",
    "priority": "优先级",
    "status": "状态",
    "closure_criteria": "闭环标准",
}
# 曾有 MINUTES_MARKERS / TITLE_MARKERS 两个常量列在这里,全代码库(含单体、脚本、
# 前端、测试)零引用。占位符识别实际走的是 PLACEHOLDER_PATTERN + ALIAS_TO_FIELD,
# 那两个名字看上去像在决定「哪些标记会被认出来」,实则什么也不决定 —— 与本轮
# 发现的 GUEST_MAX_PER_WINDOW 是同一类:定义了却从未生效的东西比没有更误导。
PLACEHOLDER_PATTERN = re.compile(
    r"\{\{\s*([^{}]+?)\s*\}\}|\$\{\s*([^{}]+?)\s*\}",
    re.IGNORECASE,
)
SEMANTIC_INSERTION_TERMS: tuple[tuple[re.Pattern[str], float, str], ...] = (
    (re.compile(r"^(会议纪要|纪要正文|完整纪要)$"), 0.92, "明确的纪要正文标题"),
    (re.compile(r"^(会议主要内容|主要内容|会议内容|访谈记录|面试记录)$"), 0.84, "可作为正文区域的章节标题"),
    (re.compile(r"(会议纪要|纪要正文|主要内容|访谈记录|面试记录)"), 0.72, "包含纪要语义的文本区域"),
)


def _template_quota_for(owner: str) -> tuple[int, int]:
    """这个账号的模板配额:(数量上限, 总量上限)。

    访客单独一档 —— 见 MAX_TEMPLATES_PER_GUEST 的说明。
    """
    if is_guest(owner):
        return MAX_TEMPLATES_PER_GUEST, MAX_TEMPLATE_STORAGE_PER_GUEST
    return MAX_TEMPLATES_PER_USER, MAX_TEMPLATE_STORAGE_PER_USER


def _quota_usage(owner: str) -> tuple[int, int]:
    """这个账号当前用掉的模板数与总字节数。"""
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        return connection.execute(
            """SELECT COALESCE(SUM(CASE WHEN deleted_at IS NULL THEN 1 ELSE 0 END), 0),
                      COALESCE(SUM(content_size), 0)
               FROM custom_templates WHERE owner_username=?""",
            (owner,),
        ).fetchone()


def _precheck_quota(owner: str, incoming_bytes: int) -> None:
    """在昂贵解析之前先拒掉必然失败的上传。

    这不是权威判定 —— 权威判定在 create_template 的写入锁内(并发上传必须由同一把
    锁串行化,否则两个请求会各自读到超限前的计数、双双放行)。这里只是让「注定被拒」
    的请求早点失败,不必先付几秒到几十秒的解析代价。
    """
    count_limit, storage_limit = _template_quota_for(owner)
    active, stored = _quota_usage(owner)
    if active >= count_limit:
        raise ValueError(f"模板数量已达上限（{count_limit} 个），请先删除不再使用的模板")
    if stored + incoming_bytes > storage_limit:
        raise ValueError(
            f"模板占用空间已达上限（{storage_limit // (1024 * 1024)}MB），"
            "已删除的模板会在 30 天后释放空间"
        )


def _clean_name(value: str, fallback: str = "自定义模板") -> str:
    # 这条正则清的是控制字符与文件名非法字符,漏了双向覆写符 —— 它们在这些范围
    # 之外,却能让模板名在列表里反向渲染,显示成另一个名字。会议标题与显示名
    # 都已经收敛过同一类字符,这里补齐最后一处。
    cleaned = BIDI_OVERRIDES.sub("", re.sub(r"[\x00-\x1f<>:\"/\\|?*]+", " ", value or "")).strip()[:60]
    return cleaned or fallback


def _normalize_alias(value: str) -> str:
    return re.sub(r"[\s_\-:：/]+", "", value or "").lower()


ALIAS_TO_FIELD = {
    _normalize_alias(alias): field
    for field, aliases in FIELD_ALIASES.items()
    for alias in aliases
}


def _canonical_field(value: str) -> str | None:
    return ALIAS_TO_FIELD.get(_normalize_alias(value))


def _validated_scenario(value: str | None, fallback: str) -> str:
    """校验场景取值,不合法直接报错。

    原先是「不在合法集合里就悄悄用回旧值/默认值」:用户在界面上改了场景、接口
    也返回成功,实际却没改,而且没有任何地方能看出问题。大小写写错(TALK)、
    误传中文标签(工程例会)、拼错(enginering)都会命中这条路径。
    同一个函数里插入方式是明确 raise 的,两者态度不应不一致。
    """
    if value is None:
        return fallback
    if value not in SCENARIOS:
        raise ValueError(f"场景取值无效：{value}")
    return value


def _migrate_columns(connection: sqlite3.Connection) -> None:
    existing = {
        str(row[1])
        for row in connection.execute("PRAGMA table_info(custom_templates)").fetchall()
    }
    columns = {
        "scenario": "TEXT NOT NULL DEFAULT 'general'",
        "is_default": "INTEGER NOT NULL DEFAULT 0",
        "parse_status": "TEXT NOT NULL DEFAULT 'ready'",
        "analysis_json": "TEXT NOT NULL DEFAULT '{}'",
        "insertion_strategy": "TEXT NOT NULL DEFAULT 'auto'",
        "insertion_target": "TEXT NOT NULL DEFAULT ''",
        "sha256": "TEXT NOT NULL DEFAULT ''",
        "content_size": "INTEGER NOT NULL DEFAULT 0",
        "version": "INTEGER NOT NULL DEFAULT 1",
        "updated_at": "REAL NOT NULL DEFAULT 0",
        "deleted_at": "REAL",
    }
    for name, definition in columns.items():
        if name not in existing:
            connection.execute(f"ALTER TABLE custom_templates ADD COLUMN {name} {definition}")

    rows = connection.execute(
        "SELECT id,content,created_at,sha256,content_size,updated_at,analysis_json "
        "FROM custom_templates"
    ).fetchall()
    for template_id, content, created_at, digest, size, updated_at, analysis_json in rows:
        raw = bytes(content)
        updates: dict[str, Any] = {}
        if not digest:
            updates["sha256"] = hashlib.sha256(raw).hexdigest()
        if not size:
            updates["content_size"] = len(raw)
        if not updated_at:
            updates["updated_at"] = created_at
        if not analysis_json or analysis_json == "{}":
            try:
                parsed = analyze_docx(raw)
                updates["analysis_json"] = json.dumps(parsed, ensure_ascii=False)
                updates["parse_status"] = parsed["parse_status"]
            except Exception as error:
                updates["analysis_json"] = json.dumps(
                    {
                        "version": 1,
                        "parse_status": "failed",
                        "risk_messages": [f"历史模板重新解析失败：{user_facing_error(error)}"],
                        "placeholders": [],
                        "structure": [],
                        "insertion_candidates": [],
                        "recommended_target": "append:new-page",
                    },
                    ensure_ascii=False,
                )
                updates["parse_status"] = "failed"
        if updates:
            assignments = ",".join(f"{name}=?" for name in updates)
            connection.execute(
                f"UPDATE custom_templates SET {assignments} WHERE id=?",
                (*updates.values(), template_id),
            )


def init_custom_template_db() -> None:
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        connection.execute(
            """CREATE TABLE IF NOT EXISTS custom_templates (
                id TEXT PRIMARY KEY,
                owner_username TEXT NOT NULL,
                name TEXT NOT NULL,
                original_filename TEXT NOT NULL,
                content BLOB NOT NULL,
                created_at REAL NOT NULL
            )"""
        )
        _migrate_columns(connection)
        connection.execute(
            "CREATE INDEX IF NOT EXISTS idx_custom_templates_owner_created "
            "ON custom_templates(owner_username, deleted_at, created_at DESC)"
        )
        duplicate_defaults = connection.execute(
            """SELECT owner_username,scenario
               FROM custom_templates
               WHERE is_default=1 AND deleted_at IS NULL
               GROUP BY owner_username,scenario
               HAVING COUNT(*) > 1"""
        ).fetchall()
        for owner, scenario in duplicate_defaults:
            rows = connection.execute(
                """SELECT id FROM custom_templates
                   WHERE owner_username=? AND scenario=? AND is_default=1 AND deleted_at IS NULL
                   ORDER BY updated_at DESC,created_at DESC,id DESC""",
                (owner, scenario),
            ).fetchall()
            for (template_id,) in rows[1:]:
                connection.execute(
                    "UPDATE custom_templates SET is_default=0 WHERE id=?",
                    (template_id,),
                )
        connection.execute(
            "CREATE UNIQUE INDEX IF NOT EXISTS idx_custom_templates_owner_default "
            "ON custom_templates(owner_username, scenario) "
            "WHERE is_default=1 AND deleted_at IS NULL"
        )


def _normalized_entry_name(item: zipfile.ZipInfo) -> str:
    """压缩包条目名的规范形式:统一分隔符、折叠 ./ 与重复斜杠、转小写。

    黑名单一律按这个结果匹配。ZIP 里同一个部件有多种等价写法
    (`word/x`、`./word/x`、`word//x`、`WORD/X`),按原样做字符串前缀比较
    只能挡住其中一种。
    """
    return str(PurePosixPath(item.filename.replace("\\", "/"))).lower()


def _check_archive_shape(infos: list[zipfile.ZipInfo]) -> None:
    """整包层面的三道:必需部件、条目数、解压后总体积。

    都只看 ZIP 头里声明的值。这是够用的:下面真正 read() 出来的字节数由
    zipfile 按 file_size 截断,声明小、实际大读不出来;声明大则在这里就被拦掉。
    """
    if not REQUIRED_DOCX_ENTRIES.issubset({item.filename for item in infos}):
        raise ValueError("文件不是有效的 DOCX 模板")
    if len(infos) > MAX_ARCHIVE_ENTRIES:
        raise ValueError("模板内部文件数量异常")
    if sum(item.file_size for item in infos) > MAX_UNCOMPRESSED_BYTES:
        raise ValueError("模板解压后体积过大")


def _check_entry(item: zipfile.ZipInfo, normalized_name: str) -> None:
    """单个条目层面的四道:路径、符号链接、压缩结构、危险部件。"""
    raw = PurePosixPath(item.filename.replace("\\", "/"))
    if raw.is_absolute() or ".." in raw.parts:
        raise ValueError("模板包含不安全路径")
    if (item.external_attr >> 16) & 0o170000 == 0o120000:
        raise ValueError("模板压缩包中不允许符号链接")
    # 先挡住 compress_size 为 0,下一行的除法才安全
    if item.file_size and not item.compress_size:
        raise ValueError("模板压缩结构异常")
    if item.file_size and item.file_size / item.compress_size > MAX_COMPRESSION_RATIO:
        raise ValueError("模板压缩比例异常")
    if any(normalized_name.startswith(prefix) for prefix in DANGEROUS_PART_PREFIXES):
        raise ValueError("模板包含宏、ActiveX 或嵌入对象")


def _check_relationships(xml_bytes: bytes) -> list[str]:
    """.rels 层面:危险关系类型、外部目标。返回需要提示用户的警告。"""
    try:
        root = ElementTree.fromstring(xml_bytes)
    except ElementTree.ParseError as error:
        raise ValueError("模板关系文件已损坏") from error
    warnings: list[str] = []
    for relationship in root:
        rel_type = relationship.attrib.get("Type", "").lower()
        if any(rel_type.endswith(suffix) for suffix in DANGEROUS_RELATIONSHIP_SUFFIXES):
            raise ValueError("模板包含外部模板或危险嵌入对象")
        if relationship.attrib.get("TargetMode", "").lower() == "external":
            # 外部超链接是正常排版的一部分,不拦但要说明;其余外部资源
            # (图片、字体、子文档)会在打开时联网取回,必须拦。
            if not rel_type.endswith("/hyperlink"):
                raise ValueError("模板包含不安全的外部资源关系")
            warnings.append("模板包含外部超链接，导出时将原样保留")
    return warnings


def _validate_archive(content: bytes) -> list[str]:
    """校验上传的 DOCX 模板压缩包,返回需要提示用户的警告。

    模板会被渲染进用户导出的 .docx 并在集团内传阅,所以这里既要挡住针对
    服务端的攻击(路径穿越、符号链接、解压炸弹、XXE),也要挡住会随文档
    传播出去的东西(宏、ActiveX、嵌入对象、外部资源关系)。

    按「整包 → 单条目 → XML 内容 → 关系」四层拆开,而不是原先一坨八层嵌套:
    每一道防线现在都能被单独读懂和单独测试。
    """
    warnings: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            infos = archive.infolist()
            _check_archive_shape(infos)
            # 所有按名字做的判断都必须用同一份规范化结果 —— 曾经穿越检查用的是
            # PurePosixPath(规范化),而危险部件的黑名单用的是只替换了反斜杠的原始
            # 字符串,于是 `./word/vbaProject.bin` 和 `word//vbaProject.bin` 两头讨好:
            # PurePosixPath 把 ./ 与 // 折叠掉,穿越检查放行;而 startswith("word/vbaproject")
            # 看到的还是带 ./ 的原样,黑名单不匹配 —— 宏就这么进来了。
            # 大小写那一维当时是想到了的(.lower()),漏的是路径这一维。
            normalized_names = {_normalized_entry_name(item) for item in infos}
            xml_total = 0
            for item in infos:
                normalized_name = _normalized_entry_name(item)
                _check_entry(item, normalized_name)
                if not normalized_name.endswith((".xml", ".rels")):
                    continue
                # 累计额度先判再读:read() 之前就把总量卡住,避免边读边超。
                xml_total += item.file_size
                if xml_total > MAX_XML_BYTES:
                    raise ValueError("模板 XML 内容体积异常")
                xml_bytes = archive.read(item)
                lowered_xml = xml_bytes.lower()
                if b"<!doctype" in lowered_xml or b"<!entity" in lowered_xml:
                    raise ValueError("模板 XML 包含不安全实体声明")
                if normalized_name.endswith(".rels"):
                    warnings.extend(_check_relationships(xml_bytes))
            content_types = archive.read("[Content_Types].xml").lower()
            if b"macroenabled" in content_types or "word/vbaproject.bin" in normalized_names:
                raise ValueError("不支持包含宏的模板")
    except zipfile.BadZipFile as error:
        raise ValueError("文件不是有效的 DOCX 模板") from error
    return list(dict.fromkeys(warnings))


def validate_docx(content: bytes, filename: str, content_type: str = "") -> list[str]:
    if not filename.lower().endswith(".docx"):
        raise ValueError("仅支持 DOCX 模板")
    if content_type.lower().split(";", 1)[0].strip() not in ALLOWED_TEMPLATE_MIME_TYPES:
        raise ValueError("上传文件类型与 DOCX 不匹配")
    if not content:
        raise ValueError("模板文件为空")
    if len(content) > MAX_TEMPLATE_BYTES:
        raise ValueError("模板文件不能超过 10 MB")
    warnings = _validate_archive(content)
    try:
        Document(io.BytesIO(content))
    except Exception as error:
        raise ValueError("DOCX 模板已损坏") from error
    return warnings


def _iter_table_paragraphs(table, prefix: str) -> Iterable[tuple[str, Paragraph, str]]:
    # Keep references to the XML nodes themselves. Using id(cell._tc) is unsafe
    # here because python-docx creates short-lived _Cell proxies and CPython may
    # reuse an id during the same traversal, causing unrelated cells to vanish
    # from analysis.
    visited_cells: set[Any] = set()
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell_element = cell._tc
            if cell_element in visited_cells:
                continue
            visited_cells.add(cell_element)
            cell_prefix = f"{prefix}/r{row_index}/c{cell_index}"
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                yield f"{cell_prefix}/p{paragraph_index}", paragraph, "table_cell"
            for table_index, nested in enumerate(cell.tables):
                yield from _iter_table_paragraphs(nested, f"{cell_prefix}/t{table_index}")


def _story_containers(document: DocumentObject) -> Iterable[tuple[str, Any]]:
    yielded: set[int] = set()
    containers: list[tuple[str, Any]] = [("body", document)]
    for section_index, section in enumerate(document.sections):
        containers.extend(
            [
                (f"header:{section_index}", section.header),
                (f"first_page_header:{section_index}", section.first_page_header),
                (f"even_page_header:{section_index}", section.even_page_header),
                (f"footer:{section_index}", section.footer),
                (f"first_page_footer:{section_index}", section.first_page_footer),
                (f"even_page_footer:{section_index}", section.even_page_footer),
            ]
        )
    for part, container in containers:
        identity = id(container._element)
        if identity in yielded:
            continue
        yielded.add(identity)
        yield part, container


def _iter_paragraphs(document: DocumentObject) -> Iterable[tuple[str, Paragraph, str, str]]:
    for part, container in _story_containers(document):
        for paragraph_index, paragraph in enumerate(container.paragraphs):
            yield f"{part}/p{paragraph_index}", paragraph, part, "paragraph"
        for table_index, table in enumerate(container.tables):
            for path, paragraph, kind in _iter_table_paragraphs(table, f"{part}/t{table_index}"):
                yield path, paragraph, part, kind


def _paragraph_style(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def _find_placeholders(text: str) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for match in PLACEHOLDER_PATTERN.finditer(text):
        alias = next(group for group in match.groups() if group is not None)
        results.append(
            {
                "raw": match.group(0),
                "alias": alias.strip(),
                "field": _canonical_field(alias),
                "start": match.start(),
                "end": match.end(),
            }
        )
    return results


def analyze_docx(content: bytes, archive_warnings: list[str] | None = None) -> dict[str, Any]:
    document = Document(io.BytesIO(content))
    placeholders: list[dict[str, Any]] = []
    structure: list[dict[str, Any]] = []
    # 两份候选分开计额度:占位符与语义标题是真正决定插入位置的,普通段落只是
    # 给人工下拉框兜底。混在一份里的话,文档开头成片的普通段落会先把名额占满,
    # 后面真正的占位符反而被丢掉 —— 推荐位置就错了。
    candidates: list[dict[str, Any]] = []
    manual_candidates: list[dict[str, Any]] = []
    truncated = False
    risks = list(archive_warnings or [])
    for path, paragraph, part, kind in _iter_paragraphs(document):
        text = paragraph.text.strip()
        if not text:
            continue
        detected = _find_placeholders(paragraph.text)
        item = {
            "path": path,
            "part": part,
            "kind": kind,
            "text": text[:500],
            "style": _paragraph_style(paragraph),
        }
        if len(structure) < MAX_PREVIEW_ITEMS:
            structure.append(item)
        for marker in detected:
            placeholder = {
                **marker,
                "path": path,
                "part": part,
                "kind": kind,
            }
            if len(placeholders) < MAX_PLACEHOLDERS:
                placeholders.append(placeholder)
            else:
                truncated = True
            if marker["field"] == "minutes" and len(candidates) < MAX_INSERTION_CANDIDATES:
                candidates.append(
                    {
                        "id": f"placeholder:{path}",
                        "path": path,
                        "placement": "replace",
                        "confidence": 1.0,
                        "label": f"替换正文占位符：{text[:80]}",
                        "reason": "模板含明确的完整纪要占位符",
                    }
                )
        semantic_candidate = False
        for pattern, confidence, reason in SEMANTIC_INSERTION_TERMS:
            if pattern.search(re.sub(r"[\s:：]+$", "", text)):
                semantic_candidate = True
                if len(candidates) >= MAX_INSERTION_CANDIDATES:
                    truncated = True
                    break
                candidates.append(
                    {
                        "id": f"after:{path}",
                        "path": path,
                        "placement": "after",
                        "confidence": confidence,
                        "label": f"插入到“{text[:80]}”之后",
                        "reason": reason,
                    }
                )
                break
        if not semantic_candidate and not any(marker["field"] == "minutes" for marker in detected):
            if len(manual_candidates) >= MAX_INSERTION_CANDIDATES:
                truncated = True
                continue
            manual_candidates.append(
                {
                    "id": f"after:{path}",
                    "path": path,
                    "placement": "after",
                    "confidence": 0.05,
                    "label": f"人工选择：插入到“{text[:80]}”之后",
                    "reason": "普通结构节点，仅在人工确认后使用",
                    "manual_only": True,
                }
            )

    candidates.extend(manual_candidates)
    if truncated:
        risks.append("模板段落过多，仅列出前若干个可选插入位置；如需精确定位请在模板中加入 {{minutes}} 占位符")
    unknown = [marker["raw"] for marker in placeholders if not marker["field"]]
    if unknown:
        risks.append(f"发现 {len(unknown)} 个未识别占位符，生成时将以“待确认”替换")
    body_markers = [marker for marker in placeholders if marker["field"] == "minutes"]
    field_markers = [marker for marker in placeholders if marker["field"]]
    candidates.append(
        {
            "id": "append:new-page",
            "path": "",
            "placement": "append",
            "confidence": 0.2,
            "label": "在保留原模板后新增页面",
            "reason": "安全兜底，不覆盖模板已有正文、签名或落款",
        }
    )
    candidates.sort(key=lambda item: (-float(item["confidence"]), item["id"]))
    recommended = candidates[0]
    if body_markers or (field_markers and not unknown):
        parse_status = "ready"
    elif recommended["confidence"] >= 0.8:
        parse_status = "ready"
    else:
        parse_status = "needs_confirmation"
        risks.append("未找到可靠的正文占位符，请确认插入位置；未确认时将在新页面追加纪要")
    return {
        "version": 2,
        "parse_status": parse_status,
        "placeholders": placeholders,
        "structure": structure,
        "insertion_candidates": candidates,
        "recommended_target": recommended["id"],
        "recommended_confidence": recommended["confidence"],
        "risk_messages": list(dict.fromkeys(risks)),
        "stats": {
            "paragraphs": len(structure),
            "placeholders": len(placeholders),
            "recognized_placeholders": len(field_markers),
            "pages_estimate": max(1, 1 + sum(1 for item in structure if "page break" in item["style"].lower())),
        },
    }


def _deserialize_analysis(raw: str) -> dict[str, Any]:
    try:
        value = json.loads(raw or "{}")
        return value if isinstance(value, dict) else {}
    except json.JSONDecodeError:
        return {}


_BASE_KEYS = (
    "id",
    "name",
    "original_filename",
    "created_at",
    "scenario",
    "is_default",
    "parse_status",
)
_TAIL_KEYS = (
    "insertion_strategy",
    "insertion_target",
    "sha256",
    "content_size",
    "version",
    "updated_at",
    "deleted_at",
)


def _row_to_template(
    row: sqlite3.Row | tuple,
    include_content: bool = False,
    *,
    with_analysis: bool = True,
) -> dict[str, Any]:
    """把数据库行映射为模板对象。

    with_analysis=False 对应 LIST_COLUMNS(列表视图):该行没有 analysis_json 列,
    analysis 置空。列表不展示解析结果,少传这一列就少了约 33KB/条。
    """
    keys = _BASE_KEYS + (("analysis_json",) if with_analysis else ()) + _TAIL_KEYS
    data = dict(zip(keys, row[: len(keys)]))
    data["filename"] = data.pop("original_filename")
    data["is_default"] = bool(data["is_default"])
    data["analysis"] = _deserialize_analysis(data.pop("analysis_json")) if with_analysis else None
    if include_content:
        data["content"] = bytes(row[len(keys)])
    return data


def template_metadata(template: dict[str, Any]) -> dict[str, Any]:
    """Return the JSON-safe, reader-facing portion of a stored template."""
    return {key: value for key, value in template.items() if key != "content"}


SELECT_COLUMNS = (
    "id,name,original_filename,created_at,scenario,is_default,parse_status,"
    "analysis_json,insertion_strategy,insertion_target,sha256,content_size,"
    "version,updated_at,deleted_at"
)
# 列表视图只展示名称、场景、大小与默认标记,不需要结构解析结果。
# analysis_json 是完整的文档大纲 + 占位符 + 插入候选,实测每条约 33KB:
# 带上它会让 40 个模板的列表响应达到 1.3MB、耗时 53ms,而其中绝大部分数据
# 界面根本不显示。详情接口(get_template)仍取全列。
LIST_COLUMNS = (
    "id,name,original_filename,created_at,scenario,is_default,parse_status,"
    "insertion_strategy,insertion_target,sha256,content_size,"
    "version,updated_at,deleted_at"
)


def create_template(
    owner: str,
    name: str,
    filename: str,
    content: bytes,
    *,
    content_type: str = "",
    scenario: str = "general",
) -> dict[str, Any]:
    # 配额先做一道廉价预检,再解析。analyze_docx 的代价随段落数线性上涨:实测
    # 8000 段要 6 秒,而 MAX_XML_BYTES(3MB)允许约五万段 —— 最坏约半分钟。
    # 原先它排在配额检查之前,于是**配额已满的账号每次被拒的上传仍要付全额解析
    # 代价**:同样是一句「模板数量已达上限」,小文件 30ms、大文件 5936ms。
    # 这条路又没有限流,等于一个免费的 CPU 消耗入口。
    #
    # 预检只查计数,不加锁 —— 它可能读到略旧的数字,但那只会往「放行」的方向偏,
    # 而真正的判定仍在下面的锁内。少拒不要紧,重点是让必然失败的请求早点失败。
    _precheck_quota(owner, len(content))
    archive_warnings = validate_docx(content, filename, content_type)
    analysis = analyze_docx(content, archive_warnings)
    clean_name = _clean_name(name, Path(filename).stem[:60] or "自定义模板")
    normalized_scenario = _validated_scenario(scenario, "general")
    template_id = uuid.uuid4().hex
    now = time.time()
    digest = hashlib.sha256(content).hexdigest()
    # 配额取值必须在拿锁之前:它内部要查 platform_users 判断是不是访客,而那同样
    # 要 PROFILE_DB_LOCK —— 这是把普通 Lock 当可重入用,会直接自锁死。
    # (第一版就写在锁里,整个测试套件卡死在这条上,栈停在 _pthread_cond_wait。)
    count_limit, storage_limit = _template_quota_for(owner)
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        # 配额检查必须和写入在同一把锁内,否则两个并发上传会各自读到超限前的
        # 计数、双双放行。PROFILE_DB_LOCK 覆盖了整段,不必再单独取写事务。
        active, stored = connection.execute(
            """SELECT COALESCE(SUM(CASE WHEN deleted_at IS NULL THEN 1 ELSE 0 END), 0),
                      COALESCE(SUM(content_size), 0)
               FROM custom_templates WHERE owner_username=?""",
            (owner,),
        ).fetchone()
        if active >= count_limit:
            raise ValueError(f"模板数量已达上限（{count_limit} 个），请先删除不再使用的模板")
        if stored + len(content) > storage_limit:
            raise ValueError(
                f"模板占用空间已达上限（{storage_limit // (1024 * 1024)}MB），"
                "已删除的模板会在 30 天后释放空间"
            )
        connection.execute(
            """INSERT INTO custom_templates(
                id,owner_username,name,original_filename,content,created_at,scenario,
                is_default,parse_status,analysis_json,insertion_strategy,insertion_target,
                sha256,content_size,version,updated_at,deleted_at
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,NULL)""",
            (
                template_id,
                owner,
                clean_name,
                Path(filename).name[:180],
                content,
                now,
                normalized_scenario,
                0,
                analysis["parse_status"],
                json.dumps(analysis, ensure_ascii=False),
                "auto",
                "",
                digest,
                len(content),
                2,
                now,
            ),
        )
    created = get_template(owner, template_id)
    return template_metadata(created) if created else {}


def list_templates(owner: str, *, include_deleted: bool = False) -> list[dict[str, Any]]:
    where = "owner_username=?" if include_deleted else "owner_username=? AND deleted_at IS NULL"
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        rows = connection.execute(
            f"SELECT {LIST_COLUMNS} FROM custom_templates WHERE {where} ORDER BY is_default DESC, updated_at DESC",
            (owner,),
        ).fetchall()
    return [_row_to_template(row, with_analysis=False) for row in rows]


def get_template(
    owner: str,
    template_id: str,
    *,
    include_deleted: bool = False,
) -> dict[str, Any] | None:
    if not template_id:
        return None
    deleted_clause = "" if include_deleted else " AND deleted_at IS NULL"
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        row = connection.execute(
            f"SELECT {SELECT_COLUMNS},content FROM custom_templates "
            f"WHERE id=? AND owner_username=?{deleted_clause}",
            (template_id, owner),
        ).fetchone()
    return _row_to_template(row, include_content=True) if row else None


def update_template(
    owner: str,
    template_id: str,
    *,
    name: str | None = None,
    scenario: str | None = None,
    is_default: bool | None = None,
    insertion_strategy: str | None = None,
    insertion_target: str | None = None,
) -> dict[str, Any] | None:
    current = get_template(owner, template_id)
    if not current:
        return None
    next_name = _clean_name(name, current["name"]) if name is not None else current["name"]
    next_scenario = _validated_scenario(scenario, current["scenario"])
    next_strategy = insertion_strategy or current["insertion_strategy"]
    if next_strategy not in {"auto", "manual", "append"}:
        raise ValueError("插入方式无效")
    valid_targets = {
        item["id"]
        for item in current["analysis"].get("insertion_candidates", [])
        if isinstance(item, dict) and item.get("id")
    }
    next_target = insertion_target if insertion_target is not None else current["insertion_target"]
    if next_strategy == "manual" and next_target not in valid_targets:
        raise ValueError("所选插入位置不存在，请重新解析模板")
    if next_strategy == "append":
        next_target = "append:new-page"
    confidence = float(current["analysis"].get("recommended_confidence") or 0)
    parse_status = (
        "ready"
        if next_strategy in {"manual", "append"} or (next_strategy == "auto" and confidence >= 0.75)
        else current["parse_status"]
    )
    default_value = current["is_default"] if is_default is None else bool(is_default)
    now = time.time()
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        connection.execute("BEGIN IMMEDIATE")
        if default_value:
            connection.execute(
                "UPDATE custom_templates SET is_default=0,updated_at=? "
                "WHERE owner_username=? AND scenario=? AND deleted_at IS NULL",
                (now, owner, next_scenario),
            )
        connection.execute(
            """UPDATE custom_templates
               SET name=?,scenario=?,is_default=?,parse_status=?,insertion_strategy=?,insertion_target=?,updated_at=?
               WHERE id=? AND owner_username=? AND deleted_at IS NULL""",
            (
                next_name,
                next_scenario,
                int(default_value),
                parse_status,
                next_strategy,
                next_target,
                now,
                template_id,
                owner,
            ),
        )
    updated = get_template(owner, template_id)
    return template_metadata(updated) if updated else None


def delete_template(owner: str, template_id: str) -> bool:
    now = time.time()
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        cursor = connection.execute(
            "UPDATE custom_templates SET deleted_at=?,is_default=0,updated_at=? "
            "WHERE id=? AND owner_username=? AND deleted_at IS NULL",
            (now, now, template_id, owner),
        )
    return cursor.rowcount > 0


# 软删除模板的内容保留期。导出接口用 include_deleted=True 读取内容 —— 历史记录
# 引用的模板即使已删除也要能重新导出,所以不能立即清空;过了保留期再释放。
DELETED_TEMPLATE_RETENTION_SECONDS = int(
    os.getenv("JKINCO_DELETED_TEMPLATE_RETENTION", str(30 * 24 * 3600))
)


def purge_templates_for_owners(owners: set[str]) -> int:
    """删除这些账号名下的模板。

    访客账号过期时,原先只清账号、资料和历史记录,模板被留了下来 ——
    实测生产已存在账号早已不存在的孤儿模板。单个模板上限 10MB,而访客是
    免注册的,不清理就是一条无人看管的存储增长路径。
    """
    if not owners:
        return 0
    marks = ",".join("?" * len(owners))
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        cursor = connection.execute(
            f"DELETE FROM custom_templates WHERE owner_username IN ({marks})", tuple(owners)
        )
    return cursor.rowcount


def purge_expired_deleted_templates(now: float | None = None) -> int:
    """释放早已软删除的模板内容。

    delete_template 只置 deleted_at,内容 BLOB 一直占着空间。保留期内仍可导出,
    过期后清空内容并置零大小,行本身保留 —— 历史记录里的模板名等元信息还要用。
    """
    cutoff = (now or time.time()) - DELETED_TEMPLATE_RETENTION_SECONDS
    with PROFILE_DB_LOCK, sqlite3.connect(PROFILE_DB) as connection:
        cursor = connection.execute(
            "UPDATE custom_templates SET content=?, content_size=0 "
            "WHERE deleted_at IS NOT NULL AND deleted_at < ? AND content_size > 0",
            (b"", cutoff),
        )
    return cursor.rowcount


def template_context(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    lines: list[str] = []
    for path, paragraph, part, _kind in _iter_paragraphs(document):
        if paragraph.text.strip():
            lines.append(f"[{part} {path}] {paragraph.text.strip()}")
    return "\n".join(lines)[:MAX_TEMPLATE_CONTEXT_CHARS]


def generate_minutes_with_template(base_minutes: str, template: dict[str, Any]) -> str:
    context = template_context(template["content"])
    prompt = f"""你是筑听平台的专业会议纪要排版助手。请严格依据已经生成并核验过的会议纪要重排内容，不得新增、猜测或虚构事实。
下面是用户上传的 DOCX 模板所体现的字段、章节、顺序和措辞。请尽量逐项对应其结构；
模板中没有信息的字段写“待确认”，不要照抄模板里的示例业务内容。

{SOURCE_GUARD}

【模板名称】{fence_source(template["name"])}
【模板结构】
{fence_source(context or "模板未包含可读取文字，请生成清晰、正式的通用会议纪要。")}

【已核验会议纪要】
{fence_source(base_minutes)}

只输出可直接写入文档的中文纪要正文，保留清晰标题和层级。"""
    return call_llm(prompt, timeout=240)


# 「显式留空」的标记,区别于「未填写」。未填写的字段要渲染成「待确认」提示用户补,
# 而纪要占位符是内容已经放到别处、这里应当什么都不留 —— 两者不能混为一谈。
BLANK_FIELD = object()


def _field_text(value: Any) -> str:
    if isinstance(value, list):
        return "\n".join(_field_text(item) for item in value)
    if isinstance(value, dict):
        return "；".join(
            f"{STRUCTURED_FIELD_LABELS.get(str(key).lower(), str(key))}：{_field_text(item)}"
            for key, item in value.items()
        )
    # 必须清控制字符。此前只有 summary/minutes 经过 clean_markdown_text,而这里是
    # 所有字段值的唯一汇聚点 —— title、participants、location、transcript,以及
    # 模型生成的 overview(映射到 {{summary}} 占位符)全都绕开了清洗。
    # 后果与当初「控制字符让 Word 导出永久失败」那次完全相同:python-docx 直接抛
    # ValueError("All strings must be XML compatible"),而 overview 由模型产出并
    # 落库 —— 里面一旦混进一个控制字符,这份纪要用自定义模板就再也导不出来。
    text = strip_control_characters(value if value is not None else "待确认").strip()
    return text if text and text.lower() not in {"undefined", "null", "none"} else "待确认"


def _replace_text_preserving_runs(paragraph: Paragraph, replacements: dict[str, Any]) -> bool:
    text = "".join(run.text for run in paragraph.runs)
    if not text:
        return False
    matches: list[tuple[int, int, str]] = []
    for match in PLACEHOLDER_PATTERN.finditer(text):
        alias = next(group for group in match.groups() if group is not None)
        field = _canonical_field(alias)
        raw_value = replacements.get(field or "", "待确认")
        value = "" if raw_value is BLANK_FIELD else _field_text(raw_value)
        matches.append((match.start(), match.end(), value))
    if not matches:
        return False
    updated = text
    for start, end, value in reversed(matches):
        updated = updated[:start] + value + updated[end:]
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def _parse_blocks(text: str) -> list[dict[str, Any]]:
    lines = clean_markdown_text(text).splitlines()
    blocks: list[dict[str, Any]] = []
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [
                [cell.strip() for cell in line.strip("|").split("|")]
                for line in table_lines
                if not re.fullmatch(r"\|?[\s:|-]+\|?", line)
            ]
            if rows:
                blocks.append({"kind": "table", "rows": rows})
            continue
        heading = re.match(r"^(#{1,6})\s*(.+)$", stripped)
        bullet = re.match(r"^[-*•]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+[.)、]\s*(.+)$", stripped)
        if heading:
            blocks.append({"kind": "heading", "level": min(3, len(heading.group(1))), "text": heading.group(2)})
        elif bullet:
            blocks.append({"kind": "bullet", "text": bullet.group(1)})
        elif numbered:
            blocks.append({"kind": "number", "text": numbered.group(1)})
        else:
            blocks.append({"kind": "paragraph", "text": stripped})
        index += 1
    return blocks or [{"kind": "paragraph", "text": "暂无可导出的结构化纪要。"}]


def _new_paragraph_after(reference: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addnext(element)
    paragraph = Paragraph(element, reference._parent)
    if style:
        try:
            paragraph.style = style
        except KeyError:
            # 用户的模板里没定义这个命名样式。回落到默认样式即可 —— 段落文字
            # 照常写入,只是少了字体字号;为此中断导出会让整份纪要都拿不到。
            pass
    paragraph.add_run(text)
    return paragraph


def _insert_blocks_after(
    document: DocumentObject,
    reference: Paragraph,
    blocks: list[dict[str, Any]],
) -> Paragraph:
    current = reference
    for block in blocks:
        kind = block["kind"]
        if kind == "table":
            rows = block["rows"]
            width = max((len(row) for row in rows), default=1)
            table = document.add_table(rows=len(rows), cols=width)
            try:
                table.style = "Table Grid"
            except KeyError:
                # 同上:模板未内置该表格样式时,表格照常生成,只是没有边框。
                pass
            for row_index, row in enumerate(rows):
                for cell_index, value in enumerate(row):
                    table.cell(row_index, cell_index).text = value
            current._p.addnext(table._tbl)
            anchor = OxmlElement("w:p")
            table._tbl.addnext(anchor)
            current = Paragraph(anchor, current._parent)
            continue
        style = None
        if kind == "heading":
            style = f"Heading {block.get('level', 1)}"
        elif kind == "bullet":
            style = "List Bullet"
        elif kind == "number":
            style = "List Number"
        current = _new_paragraph_after(current, str(block.get("text", "")), style)
    return current


def _blocks_for_field(field: str, value: Any) -> list[dict[str, Any]]:
    if isinstance(value, list):
        if value and all(isinstance(item, dict) for item in value):
            keys: list[str] = []
            for item in value:
                for key in item:
                    if key not in keys:
                        keys.append(str(key))
            rows = [[STRUCTURED_FIELD_LABELS.get(key.lower(), key) for key in keys]]
            rows.extend([[_field_text(item.get(key, "")) for key in keys] for item in value])
            return [{"kind": "table", "rows": rows}]
        return [{"kind": "bullet", "text": _field_text(item)} for item in value]
    parsed = _parse_blocks(_field_text(value))
    if field in LIST_FIELDS and all(block["kind"] == "paragraph" for block in parsed):
        return [{"kind": "bullet", "text": block["text"]} for block in parsed]
    return parsed


def _field_values(summary: str, fields: dict[str, Any] | None = None) -> dict[str, Any]:
    cleaned = clean_markdown_text(summary)
    source = dict(fields or {})
    title = str(source.get("title") or "").strip()
    if not title:
        title = next(
            (line.lstrip("# ").strip() for line in cleaned.splitlines() if line.strip()),
            "会议纪要",
        )
    defaults: dict[str, Any] = {
        "title": title,
        "meeting_type": source.get("meeting_type") or source.get("mode_label") or "通用会议纪要",
        "date": source.get("date") or time.strftime("%Y-%m-%d"),
        "start_time": source.get("start_time") or "待确认",
        "end_time": source.get("end_time") or "待确认",
        "location": source.get("location") or "待确认",
        "participants": source.get("participants") or "待确认",
        "host": source.get("host") or "待确认",
        "recorder": source.get("recorder") or "待确认",
        "client": source.get("client") or "待确认",
        "project": source.get("project") or "待确认",
        "summary": source.get("overview") or source.get("summary") or cleaned[:600],
        "minutes": cleaned,
        "agenda": source.get("agenda") or cleaned,
        "conclusions": source.get("conclusions") or cleaned,
        "todos": source.get("todos") or cleaned,
        "transcript": source.get("transcript") or "未提供原始转写",
        "generated_at": source.get("generated_at") or time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    result: dict[str, Any] = {}
    for key, value in defaults.items():
        result[key] = value if isinstance(value, (list, dict)) else _field_text(value)
    return result


def _resolve_target(
    analysis: dict[str, Any],
    insertion_strategy: str,
    insertion_target: str,
) -> str:
    if insertion_strategy == "append":
        return "append:new-page"
    if insertion_strategy == "manual" and insertion_target:
        return insertion_target
    recommended = str(analysis.get("recommended_target") or "")
    confidence = float(analysis.get("recommended_confidence") or 0)
    return recommended if confidence >= 0.75 else "append:new-page"


def _render_into_document(
    document: DocumentObject,
    summary: str,
    values: dict[str, Any],
    analysis: dict[str, Any],
    insertion_strategy: str,
    insertion_target: str,
) -> None:
    path_map = {path: paragraph for path, paragraph, _part, _kind in _iter_paragraphs(document)}
    target = _resolve_target(analysis, insertion_strategy, insertion_target)
    minutes_inserted = False
    target_path = target.split(":", 1)[1] if ":" in target else ""

    if target.startswith("placeholder:") and target_path in path_map:
        paragraph = path_map[target_path]
        text = "".join(run.text for run in paragraph.runs) or paragraph.text
        marker_match = next(
            (
                match
                for match in PLACEHOLDER_PATTERN.finditer(text)
                if _canonical_field(next(group for group in match.groups() if group is not None)) == "minutes"
            ),
            None,
        )
        if marker_match:
            prefix = text[: marker_match.start()].strip()
            suffix = text[marker_match.end() :].strip()
            if paragraph.runs:
                paragraph.runs[0].text = prefix
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(prefix)
            blocks = _parse_blocks(summary)
            last = _insert_blocks_after(document, paragraph, blocks)
            if suffix:
                _new_paragraph_after(last, suffix)
            minutes_inserted = True
    elif target.startswith("after:") and target_path in path_map:
        _insert_blocks_after(document, path_map[target_path], _parse_blocks(summary))
        minutes_inserted = True

    structured_elements: set[int] = set()
    for _path, paragraph, _part, _kind in list(_iter_paragraphs(document)):
        text = "".join(run.text for run in paragraph.runs) or paragraph.text
        markers = _find_placeholders(text)
        if len(markers) != 1:
            continue
        marker = markers[0]
        field = marker["field"]
        if field not in LIST_FIELDS or text.strip() != marker["raw"].strip():
            continue
        if paragraph.runs:
            paragraph.runs[0].text = ""
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run("")
        _insert_blocks_after(document, paragraph, _blocks_for_field(field, values[field]))
        structured_elements.add(id(paragraph._element))

    scalar_values = {key: value for key, value in values.items() if key != "minutes"}
    for _path, paragraph, _part, _kind in list(_iter_paragraphs(document)):
        if id(paragraph._element) in structured_elements:
            continue
        markers = _find_placeholders("".join(run.text for run in paragraph.runs) or paragraph.text)
        if not markers:
            continue
        if any(marker["field"] == "minutes" for marker in markers):
            # 占位符一律清空。纪要正文要么已按块插到占位符处(minutes_inserted),
            # 要么会由下面的 append 分支追加 —— 两条路必有其一,不存在「都没插」的
            # 情况。原先在未插入时把原始 markdown 填进占位符,结果是纪要出现两遍:
            # 一遍带着 ## 和 - 符号,一遍是格式化的。选「追加到新页」的模板必中。
            scalar_values["minutes"] = BLANK_FIELD
        _replace_text_preserving_runs(paragraph, scalar_values)

    if not minutes_inserted:
        document.add_page_break()
        anchor = document.add_paragraph()
        _insert_blocks_after(document, anchor, _parse_blocks(summary))
        # The anchor only establishes a stable insertion point after the page break.
        anchor._element.getparent().remove(anchor._element)

    for _path, paragraph, _part, _kind in list(_iter_paragraphs(document)):
        _replace_text_preserving_runs(paragraph, values)


def render_custom_docx(
    content: bytes,
    summary: str,
    output: Path,
    *,
    fields: dict[str, Any] | None = None,
    analysis: dict[str, Any] | None = None,
    insertion_strategy: str = "auto",
    insertion_target: str = "",
) -> Path:
    document = Document(io.BytesIO(content))
    values = _field_values(summary, fields)
    current_analysis = analysis or analyze_docx(content)
    _render_into_document(
        document,
        summary,
        values,
        current_analysis,
        insertion_strategy,
        insertion_target,
    )
    # docx 的核心属性有 255 字符硬上限,超了 python-docx 直接抛 ValueError。
    # 调用方传的 title 有长度约束,但没传时会从纪要首行推导(见 _field_values),
    # 而那一行不受任何约束 —— 首行一长,整份导出就失败。属性只是文件元数据,
    # 截断无损;正文里的 {{title}} 占位符仍用完整文本。
    document.core_properties.title = _field_text(values["title"])[:255]
    document.core_properties.subject = "筑听平台自定义模板会议纪要"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
