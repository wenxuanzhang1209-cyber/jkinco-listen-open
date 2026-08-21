"""Source-faithful and branded report builders for JKinco Listen."""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


MODE_META = {
    "general": {
        "title": "会议纪要",
        "subtitle": "议题、决议、行动与待确认事项",
        "accent": "2467DC",
        "chart": ("议题", "决议", "待办", "风险"),
    },
    "lingxi": {
        "title": "会议纪要",
        "subtitle": "议题、决议、行动与待确认事项",
        "accent": "2467DC",
        "chart": ("议题", "决议", "待办", "风险"),
    },
    "personal": {
        "title": "个人备忘录",
        "subtitle": "工作进展、思考判断与行动计划",
        "accent": "0F8F83",
        "chart": ("工作进展", "待办", "风险", "明日重点"),
    },
    "interview": {
        "title": "面试记录与候选人反馈表",
        "subtitle": "事实记录、能力证据与录用建议",
        "accent": "6B5CC5",
        "chart": ("能力证据", "优势", "风险", "后续安排"),
    },
    "customer_visit": {
        "title": "客户拜访会议纪要",
        "subtitle": "客户诉求、沟通结论与商机跟进",
        "accent": "C96D20",
        "chart": ("客户诉求", "沟通要点", "待办", "风险"),
    },
    "auto": {
        "title": "智能识别场景路由",
        "subtitle": "自动区分工程例会与通用会议纪要",
        "accent": "2467DC",
        "chart": ("工程例会", "通用纪要", "个人备忘", "业务记录"),
    },
}


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def _cell_margins(cell, top=100, start=130, bottom=100, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_run(run, size=10.5, bold=False, color="172437", font="宋体") -> None:
    run.font.name = font
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    fonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_paragraph_text(paragraph, text: str, *, size=10.5, bold=False, color="172437", font="宋体") -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    _set_run(run, size=size, bold=bold, color=color, font=font)


def _set_cell_text(cell, text: str, *, bold=False, align=None, size=10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    _set_paragraph_text(paragraph, text, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_margins(cell)


def _set_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def _remove_body_images(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        embeds = paragraph._p.xpath(".//a:blip/@r:embed")
        if not embeds:
            continue
        for rel_id in embeds:
            try:
                document.part.drop_rel(rel_id)
            except KeyError:
                pass
        paragraph._element.getparent().remove(paragraph._element)


HEADER_COMPANY_NAME = "上海建科工程咨询有限公司"


def _rebrand_header(document: Document, company: str = HEADER_COMPANY_NAME) -> None:
    """Replace the source company name in every header while keeping run styling."""
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if "思立博" not in paragraph.text and company not in paragraph.text:
                continue
            replaced = False
            for run in paragraph.runs:
                if "思立博" in run.text:
                    run.text = company
                    replaced = True
                elif replaced and run.text.strip():
                    run.text = ""
            if not replaced and paragraph.runs:
                paragraph.runs[0].text = company


def prepare_talk_template(reference: Path, output: Path, logo_output: Path | None = None) -> Path:
    """Create a clean template while preserving the source layout and header logo."""
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(reference, output)
    document = Document(output)
    if len(document.tables) < 2 or len(document.tables[0].rows) < 6:
        raise ValueError("工地例会模板结构与预期不一致")

    _remove_body_images(document)
    _rebrand_header(document)
    _set_paragraph_text(document.paragraphs[1], "  会 议 纪 要", size=18, bold=True, font="黑体")
    _set_paragraph_text(document.paragraphs[2], "项目名称：{{PROJECT_NAME}}                         编 号：{{MEETING_NO}}", size=10.5)

    table = document.tables[0]
    values = ("{{MEETING_NAME}}", "{{MEETING_TIME}}", "{{MEETING_LOCATION}}", "{{ATTENDEES}}", "{{HOST}}")
    for row_index, value in enumerate(values):
        _set_cell_text(table.cell(row_index, 1), value, align=WD_ALIGN_PARAGRAPH.CENTER)
    content_cell = table.cell(5, 0)
    content_cell.text = ""
    _set_paragraph_text(content_cell.paragraphs[0], "会议内容：", bold=True, font="黑体")
    marker = content_cell.add_paragraph()
    _set_paragraph_text(marker, "{{REPORT_BODY}}")

    sign_table = document.tables[1]
    for index, label in enumerate(("施工单位", "监理单位", "建设单位")):
        _set_cell_text(sign_table.cell(0, index), label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    document.core_properties.title = "筑听平台工程会议纪要模板"
    document.core_properties.subject = "空白工程例会会议纪要模板"
    document.core_properties.author = "上海建科咨询集团"
    document.save(output)

    if logo_output:
        logo_output.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(reference) as archive:
            logo_output.write_bytes(archive.read("word/media/image1.jpeg"))
    return output


def _clean_lines(text: str) -> list[str]:
    clean = str(text or "").replace("```markdown", "").replace("```", "")
    clean = re.sub(r"\*\*(.*?)\*\*", r"\1", clean)
    return [line.strip() for line in clean.splitlines()]


def _extract_value(text: str, labels: tuple[str, ...], fallback="未提及") -> str:
    for line in _clean_lines(text):
        normalized = re.sub(r"^[#\-*•\d.、\s]+", "", line)
        for label in labels:
            match = re.match(rf"^{re.escape(label)}\s*[：:]\s*(.+)$", normalized)
            if match and match.group(1).strip():
                return match.group(1).strip()
    return fallback


def _report_body_lines(summary: str) -> list[str]:
    skip = ("项目名称", "编号", "会议名称", "会议时间", "会议地点", "与会人员", "主持人")
    lines: list[str] = []
    for line in _clean_lines(summary):
        plain = line.lstrip("#").strip()
        if not plain or plain.replace(" ", "") in {"会议纪要", "会议内容：", "会议内容"}:
            continue
        if any(re.match(rf"^{label}\s*[：:]", plain) for label in skip):
            continue
        lines.append(plain)
    return lines or ["一、会议内容", "未提及"]


def _add_talk_line(cell, line: str) -> None:
    paragraph = cell.add_paragraph()
    major = bool(re.match(r"^[一二三四五六七八九十]+、", line))
    subsection = bool(re.match(r"^[（(]\d+[）)]", line)) or (
        len(line) <= 22 and line.endswith(("情况", "管理", "事项", "计划", "要求"))
    )
    paragraph.paragraph_format.space_before = Pt(7 if major else 3 if subsection else 0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.2
    _set_paragraph_text(
        paragraph,
        line,
        size=10.5,
        bold=major or subsection,
        font="黑体" if major or subsection else "宋体",
    )


def build_talk_report(summary: str, template: Path, output: Path) -> Path:
    document = Document(template)
    _rebrand_header(document)
    project_name = _extract_value(summary, ("项目名称",))
    meeting_no = _extract_value(summary, ("编号", "会议编号"), "待确认")
    _set_paragraph_text(
        document.paragraphs[2],
        f"项目名称：{project_name}                         编 号：{meeting_no}",
        size=10.5,
    )
    table = document.tables[0]
    metadata = (
        _extract_value(summary, ("会议名称", "会议主题")),
        _extract_value(summary, ("会议时间", "时间")),
        _extract_value(summary, ("会议地点", "地点")),
        _extract_value(summary, ("与会人员", "参会人员")),
        _extract_value(summary, ("主持人",)),
    )
    for row_index, value in enumerate(metadata):
        _set_cell_text(table.cell(row_index, 1), value, align=WD_ALIGN_PARAGRAPH.CENTER)

    content_cell = table.cell(5, 0)
    content_cell.text = ""
    _set_paragraph_text(content_cell.paragraphs[0], "会议内容：", bold=True, font="黑体")
    for line in _report_body_lines(summary):
        _add_talk_line(content_cell, line)

    document.core_properties.title = metadata[0]
    document.core_properties.subject = "筑听平台自动生成工程会议纪要"
    document.core_properties.author = "上海建科咨询集团"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def _add_bottom_border(paragraph, color="1F6B8A", size="10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    for kind, value in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        else:
            node = OxmlElement("w:instrText" if value == " PAGE " else "w:t")
            node.text = value
        run._r.append(node)


def _font_path() -> str | None:
    candidates = (
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    )
    return next((path for path in candidates if Path(path).exists()), None)


def _chart_png(labels: tuple[str, ...], values: list[int], accent: str, output: Path) -> None:
    width, height = 1280, 360
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = _font_path()
    font = ImageFont.truetype(font_path, 28) if font_path else ImageFont.load_default()
    small = ImageFont.truetype(font_path, 23) if font_path else ImageFont.load_default()
    rgb = tuple(int(accent[index:index + 2], 16) for index in (0, 2, 4))
    max_value = max(max(values), 1)
    for index, (label, value) in enumerate(zip(labels, values)):
        y = 48 + index * 72
        draw.text((30, y), label, fill=(48, 65, 86), font=font)
        draw.rounded_rectangle((235, y + 5, 1120, y + 37), radius=16, fill=(232, 238, 247))
        bar_width = max(24, int(885 * value / max_value)) if value else 8
        draw.rounded_rectangle((235, y + 5, 235 + bar_width, y + 37), radius=16, fill=rgb)
        draw.text((1150, y), str(value), fill=(48, 65, 86), font=small)
    image.save(output)


def _metric_values(text: str, labels: tuple[str, ...]) -> list[int]:
    clean = str(text or "")
    synonyms = {
        "重点事项": ("重点", "关键", "重要"), "待办": ("待办", "下一步", "后续"),
        "风险": ("风险", "隐患", "阻塞", "关注"), "决策": ("决策", "协调", "决定"),
        "工作进展": ("进展", "完成", "推进"), "明日重点": ("明日", "明天", "优先"),
        "能力证据": ("能力", "评价", "依据"), "优势": ("优势", "匹配", "亮点"),
        "后续安排": ("后续", "复面", "安排"), "客户诉求": ("诉求", "需求", "关注"),
        "沟通要点": ("沟通", "反馈", "回应"), "个人备忘": ("备忘", "复盘", "想法"),
        "会议纪要": ("施工", "监理", "会议"), "管理简报": ("管理", "汇报", "团队"),
        "业务记录": ("客户", "面试", "岗位"),
    }
    return [sum(clean.count(term) for term in synonyms.get(label, (label,))) for label in labels]


def _add_pipe_table(document: Document, rows: list[str], accent: str) -> None:
    values = [[item.strip() or "待确认" for item in row.strip().strip("|").split("|")] for row in rows]
    if len(values) > 1 and all(re.fullmatch(r"[-: ]+", item) for item in values[1]):
        values.pop(1)
    columns = max(len(row) for row in values)
    table = document.add_table(rows=len(values), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for row_index, row in enumerate(values):
        _set_row_no_split(table.rows[row_index])
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            value = row[column_index] if column_index < len(row) else ""
            _set_cell_text(cell, value, bold=row_index == 0, size=9.5)
            if row_index == 0:
                _shade(cell, accent)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_report_content(document: Document, summary: str, accent: str, title: str) -> None:
    table_buffer: list[str] = []
    for raw in _clean_lines(summary):
        line = raw.strip()
        if not line or line.replace(" ", "") == title.replace(" ", ""):
            if table_buffer:
                _add_pipe_table(document, table_buffer, accent)
                table_buffer = []
            continue
        if "|" in line and not line.startswith("#"):
            table_buffer.append(line)
            continue
        if table_buffer:
            _add_pipe_table(document, table_buffer, accent)
            table_buffer = []
        heading = line.startswith("#") or bool(re.match(r"^[一二三四五六七八九十]+、", line))
        if heading:
            text = line.lstrip("#").strip()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.keep_with_next = True
            _set_paragraph_text(paragraph, text, size=13, bold=True, color=accent, font="黑体")
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Mm(4) if re.match(r"^(?:[-*•]|\d+[.、])", line) else Mm(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.3
        _set_paragraph_text(paragraph, line, size=10.5)
    if table_buffer:
        _add_pipe_table(document, table_buffer, accent)


def build_branded_report(summary: str, mode: str, logo: Path, output: Path) -> Path:
    meta = MODE_META.get(mode, MODE_META["auto"])
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(19)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    header = section.header.paragraphs[0]
    if logo.exists():
        header.add_run().add_picture(str(logo), width=Mm(15))
    _add_bottom_border(header, meta["accent"], "8")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(footer.add_run("筑听平台 · 上海建科咨询集团  |  "), size=8.5, color="7A8AA2")
    _add_page_field(footer)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    _set_paragraph_text(title, meta["title"], size=22, bold=True, color=meta["accent"], font="黑体")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_paragraph_text(subtitle, meta["subtitle"], size=10.5, color="6B7C93")

    metrics = _metric_values(summary, meta["chart"])
    dashboard = document.add_table(rows=1, cols=4)
    dashboard.alignment = WD_TABLE_ALIGNMENT.CENTER
    dashboard.autofit = False
    for index, (label, value) in enumerate(zip(meta["chart"], metrics)):
        cell = dashboard.cell(0, index)
        _shade(cell, "F2F6FC")
        _cell_margins(cell, 130, 130, 130, 130)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_text(paragraph, label, size=8.5, color="6B7C93")
        value_p = cell.add_paragraph()
        value_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_text(value_p, str(value), size=17, bold=True, color=meta["accent"], font="黑体")
    _set_row_no_split(dashboard.rows[0])

    chart_heading = document.add_paragraph()
    chart_heading.paragraph_format.space_before = Pt(12)
    chart_heading.paragraph_format.space_after = Pt(3)
    _set_paragraph_text(chart_heading, "智能信息概览", size=12, bold=True, color=meta["accent"], font="黑体")
    with tempfile.TemporaryDirectory(prefix="jkinco-chart-") as temp_dir:
        chart_path = Path(temp_dir) / "chart.png"
        _chart_png(meta["chart"], metrics, meta["accent"], chart_path)
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(chart_path), width=Mm(165))

    _add_report_content(document, summary, meta["accent"], meta["title"])
    document.core_properties.title = meta["title"]
    document.core_properties.subject = "筑听平台场景化结构化报告"
    document.core_properties.author = "上海建科咨询集团"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def build_report_docx(summary: str, mode: str, output: Path, *, talk_template: Path, logo: Path) -> Path:
    normalized = "talk" if mode in {"talk", "engineering"} else mode
    if normalized == "talk":
        return build_talk_report(summary, talk_template, output)
    return build_branded_report(summary, normalized, logo, output)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    bundled = Path.home() / ".cache/codex-runtimes/codex-primary-runtime/dependencies/bin/override/soffice"
    executable = str(bundled) if bundled.exists() else shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        return False
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="jkinco-lo-") as profile:
        command = [
            executable,
            f"-env:UserInstallation=file://{profile}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ]
        environment = os.environ.copy()
        font_config = Path(__file__).resolve().parent / "docs" / "fontconfig-macos.xml"
        if font_config.exists():
            environment["FONTCONFIG_FILE"] = str(font_config)
        completed = subprocess.run(command, capture_output=True, text=True, timeout=120, env=environment)
        generated = pdf_path.parent / f"{docx_path.stem}.pdf"
        if completed.returncode != 0 or not generated.exists():
            return False
        if generated != pdf_path:
            generated.replace(pdf_path)
    return pdf_path.exists() and pdf_path.stat().st_size > 0
