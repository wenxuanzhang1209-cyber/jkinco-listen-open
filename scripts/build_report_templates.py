#!/usr/bin/env python3
"""Build the four clean production report templates."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "templates"

TEMPLATES = [
    (
        "筑言工程监理会议纪要模板.docx",
        "会 议 纪 要",
        "筑言｜工程监理会议标准输出",
        "工程例会 · 质量安全 · 进度验收 · 责任闭环",
        "0B4F8A",
    ),
    (
        "个人备忘录总结_会议纪要模板.docx",
        "个人备忘录总结会议纪要",
        "个人助手｜语音备忘与工作复盘",
        "摘要结论 · 记录要点 · 待办事项 · 风险提醒 · 个人复盘",
        "087F8C",
    ),
    (
        "HR面试记录与候选人反馈模板.docx",
        "面试记录与候选人反馈表",
        "HR｜面试过程记录与人才决策",
        "基础信息 · 过程记录 · 能力评价 · 风险关注 · 录用建议",
        "2F5D8C",
    ),
    (
        "客户拜访会议纪要模板_建科.docx",
        "客户拜访会议纪要",
        "建科｜客户需求与合作跟进",
        "拜访背景 · 客户诉求 · 沟通要点 · 责任分工 · 后续计划",
        "123F73",
    ),
]


def set_font(run, size, color="172437", bold=False, font="Hiragino Sans GB"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), font)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def add_bottom_border(paragraph, color):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def build_template(filename, title_text, subtitle_text, scope_text, accent):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.71)
    section.bottom_margin = Inches(0.63)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Hiragino Sans GB"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Hiragino Sans GB")

    for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        style = doc.styles[name]
        style.font.name = "Hiragino Sans GB"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(accent)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Hiragino Sans GB")

    header = section.header.paragraphs[0]
    set_font(header.add_run("JKINCO  |  筑听 · 筑言智能报告"), 8.5, "6B7C93", True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("内部工作文档  |  第 "), 8.5, "8090A5")
    add_page_field(footer)
    set_font(footer.add_run(" 页"), 8.5, "8090A5")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run(title_text), 20, accent, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    set_font(subtitle.add_run(subtitle_text), 10.5, "53677F", True)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.space_after = Pt(14)
    set_font(scope.add_run(scope_text), 9, "728399")
    add_bottom_border(scope, accent)

    marker = doc.add_paragraph("{{REPORT_BODY}}")
    marker.paragraph_format.space_before = Pt(4)
    marker.paragraph_format.space_after = Pt(0)
    marker.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.core_properties.title = title_text
    doc.core_properties.subject = "筑听/筑言四大场景生产模板"
    doc.core_properties.author = "JKinco"
    doc.save(TEMPLATE_DIR / filename)


def main():
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    for args in TEMPLATES:
        build_template(*args)
        print(TEMPLATE_DIR / args[0])


if __name__ == "__main__":
    main()
