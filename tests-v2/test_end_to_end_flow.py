"""整条链跑一遍:注册 → 建会 → 加入 → 聊天 → 转写 → 出纪要 → 导出 → 历史。

单个护栏各自有用例,但它们是分批加上去的,彼此之间没有人验证过。这一轮就加了
十来处收敛(标题、显示名、聊天、异常文案、空转写闸门、验证码、纪要写权限、
头像尺寸),每一处都在别人的输入路径上 —— 叠在一起最容易出的问题是「后一道把
前一道需要的东西改没了」,而那种问题在单元用例里看不见。

这条用例故意在每个入口都塞进要被收敛的内容,最后检查:
  - 该清掉的都清掉了(控制字符、双向覆写符、伪造的发言行);
  - 该留下的一样没少(emoji、正常姓名、多行消息、纪要正文)。
"""
from __future__ import annotations

import io
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient

import backend.meetings as meetings
from backend.main import app
from helpers import solve_captcha

MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ADMIN_USERNAME, _, ADMIN_PASSWORD = meetings.os.environ["JKINCO_AUTH"].split(",", 1)[0].partition(":")


def _template_bytes() -> bytes:
    document = Document()
    for text in ("会议：{{title}}", "参会：{{participants}}", "概览：{{summary}}", "{{minutes}}"):
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_the_whole_flow_survives_every_guard():
    with TestClient(app) as host, TestClient(app) as guest:
        # 1) 注册:验证码答案不再能从 token 里读出来,只能真的解
        challenge = guest.get("/api/auth/captcha").json()
        registered = guest.post("/api/auth/register", json={
            "username": "e2e_member", "display_name": "李四\n王五：我反对",   # 伪造署名的尝试
            "password": "StrongPass123",
            "captcha_token": challenge["token"], "captcha_answer": solve_captcha(challenge["token"]),
        })
        assert registered.status_code == 201, registered.text
        assert "\n" not in registered.json()["user"]["display_name"], "显示名的换行没被压平"

        # 2) 建会:标题带控制字符与双向覆写符,emoji 要留下
        host.post("/api/auth/login", json={"username": ADMIN_USERNAME, "password": ADMIN_PASSWORD})
        meeting = host.post("/api/meetings", json={"title": "三标段周会 🎉\x00abc‮gfd"}).json()
        assert not any(char in meeting["title"] for char in "\x00‮"), meeting["title"]
        assert "🎉" in meeting["title"], "emoji 被误删"

        code, meeting_id = meeting["meeting_code"], meeting["id"]
        host.post(f"/api/meetings/{code}/join", json={"display_name": "主持人"})
        joined = guest.post(f"/api/meetings/{code}/join", json={"display_name": "李四\n王五：我反对"})
        assert joined.status_code == 200, joined.text

        # 3) 聊天:控制字符清掉,多行保留
        sent = guest.post(f"/api/meetings/{meeting_id}/chat", json={"message": "第一行\n第二行\x00"})
        assert sent.json()["message"] == "第一行\n第二行"

        # 4) 转写:参会者说一句,伪造的署名不得凭空造出一行
        meetings._store_transcript(meeting_id, joined.json()["identity"], {
            "text": "我这边进度正常", "sentence_end": True, "sentence_id": 1,
            "begin_time": 0, "end_time": 1000,
        })
        with meetings.db() as connection:
            transcript = meetings._final_transcript_text(connection, meeting_id)
        lines = [line for line in transcript.splitlines() if line.strip()]
        assert len(lines) == 1, f"署名伪造出了额外的发言行:{lines}"
        assert lines[0].count("：") == 1, f"说话人有歧义:{lines[0]!r}"

        # 5) 参会者不得改写正式纪要
        assert guest.patch(f"/api/meetings/{meeting_id}/minutes",
                           json={"content_markdown": "## 伪造结论"}).status_code == 403

        # 6) 出纪要:主持人定稿
        assert host.patch(f"/api/meetings/{meeting_id}/minutes",
                          json={"content_markdown": "## 一、结论\n- 方案通过 🎉"}).status_code == 200

        # 7) 用自定义模板导出:概览里带控制字符也不能让导出失败
        created = host.post("/api/custom-templates", files={"file": ("t.docx", _template_bytes(), MIME)},
                            data={"name": "周会模板 🎉", "scenario": "general"})
        assert created.status_code == 201, created.text
        exported = host.post("/api/export/docx", json={
            "summary": "## 一、结论\n- 方案通过 🎉",
            "overview": "概览\x0b要点",              # 模型产出可能带控制字符
            "participants": "张三、李四 🎉",
            "mode": "general",
            "custom_template_id": created.json()["id"],
        })
        assert exported.status_code == 200, exported.text[:200]
        body = "\n".join(p.text for p in Document(io.BytesIO(exported.content)).paragraphs)
        assert "方案通过" in body and "🎉" in body, "正文内容丢失"
        assert "\x0b" not in body


def test_a_silent_meeting_produces_no_fabricated_minutes():
    """没人说话的那条路同样要走通:标记为「无有效发言」,而不是造一份纪要。"""
    with TestClient(app) as client:
        client.post("/api/auth/login", json={"username": ADMIN_USERNAME, "password": ADMIN_PASSWORD})
        meeting = client.post("/api/meetings", json={"title": "没人说话的会"}).json()
        client.post(f"/api/meetings/{meeting['meeting_code']}/join", json={"display_name": "主持人"})
        client.post(f"/api/meetings/{meeting['id']}/start")

        joined = client.get(f"/api/meetings/{meeting['id']}").json()
        assert joined["status"] == "active"
        # 只有噪音
        meetings._store_transcript(meeting["id"], client.post(
            f"/api/meetings/{meeting['meeting_code']}/join", json={"display_name": "主持人"}
        ).json()["identity"], {"text": "嗯。", "sentence_end": True, "sentence_id": 7,
                              "begin_time": 0, "end_time": 100})

        with patch.object(meetings.core, "generate_minutes", side_effect=AssertionError("不该生成纪要")):
            meetings._finalize_minutes(meeting["id"])

        detail = client.get(f"/api/meetings/{meeting['id']}").json()
        assert detail["minutes_status"] == "empty", detail["minutes_status"]
