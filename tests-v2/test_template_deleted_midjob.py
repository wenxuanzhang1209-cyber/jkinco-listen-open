"""任务运行期间模板被删,不能把已经转写好的内容一起丢掉。

run_processing_job 里有一段明确的注释(「转写到这里已经完成 —— 一小时录音要跑
几十秒、六小时要几分钟……后续任何一步失败,都不能让转写结果一起丢掉:实时录音
根本无法重录」),下面用 try/except 把纪要生成的失败兜住了,失败也照样落库。

但自定义模板的存在性检查排在那道保护**之前**:

    custom_template = get_template(username, custom_template_id) ...
    if custom_template_id and not custom_template:
        raise RuntimeError("所选自定义模板不存在或无权使用")   # ← 落到外层 except

提交时(/api/process)已经校验过模板存在,所以这条路只有一种走法:模板是在
**转写进行当中**被删掉的。删除是软删(deleted_at),而 get_template 默认过滤掉
已删除的记录 —— 于是一小时的转写就这么没了。

佐证这是疏漏而不是有意为之:导出那条路径早就写成了 include_deleted=True
(main.py 的 export),也就是「用过的模板事后被删,不该让流程失败」这个判断
作者已经做过一次,只是处理流水线这条漏了。
"""
from __future__ import annotations

import os
import tempfile
from unittest.mock import patch

os.environ.setdefault("JKINCO_HISTORY_DIR", tempfile.mkdtemp(prefix="jkinco-tpl-race-"))

import backend.main as main
from backend.custom_templates import create_template, delete_template, get_template

OWNER = "tester"


def _docx_bytes() -> bytes:
    """现造一份最小但合法的 docx —— 上传接口会真的解析它。"""
    import io

    from docx import Document

    document = Document()
    document.add_heading("会议纪要", level=1)
    document.add_paragraph("一、会议概况")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_template() -> str:
    template = create_template(
        OWNER,
        name="例会模板",
        filename="t.docx",
        content=_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        scenario="talk",
    )
    return template["id"]


def _run_job_with(template_id: str) -> dict:
    """跑一次任务,返回落库时拿到的参数(没落库则为空)。"""
    saved: dict = {}

    def fake_save(transcript, summary, status, mode, source, overview, owner_username="", **_metadata):
        saved.update(transcript=transcript, summary=summary, status=status)
        return "rec-1"

    with patch.object(main.core, "infer_app_mode_best_effort", return_value=("talk", "理由")), \
         patch.object(main.core, "mode_label", return_value="工程例会"), \
         patch.object(main.core, "should_generate_and_push", return_value=False), \
         patch.object(main.core, "should_push_to_dingtalk", return_value=False), \
         patch.object(main.core, "save_meeting_history_record", side_effect=fake_save), \
         patch.object(main, "set_job"):
        main.JOB_CAPACITY.acquire()
        main.run_processing_job(
            "job-race", None, "今天例会讨论了监理单位的旁站安排和检验批验收",
            "只转写，不推送", "auto", OWNER, template_id, "live",
        )
    return saved


def test_transcript_survives_when_template_is_deleted_during_the_job():
    template_id = _make_template()
    # 提交时模板还在 —— 这正是 /api/process 校验通过的那一刻
    assert get_template(OWNER, template_id) is not None
    # 转写进行中用户把它删了
    assert delete_template(OWNER, template_id) is True

    saved = _run_job_with(template_id)

    assert saved, "模板被删导致整段转写被丢弃 —— 实时录音无法重录，这是不可挽回的数据损失"
    assert "监理单位" in saved["transcript"], f"落库的转写不对: {saved.get('transcript')!r}"


def test_a_template_that_never_existed_still_fails_loudly():
    """反向:模板 id 根本是伪造的,不能默默当成没选模板。

    上面那条的修法如果做成「查不到就忽略」,就会把这种情况也一起放过 ——
    用户会拿到一份没套模板的纪要,而界面上显示的是他选的那个模板。
    """
    saved = _run_job_with("does-not-exist-at-all")
    assert not saved, "伪造的模板 id 被静默忽略了"
