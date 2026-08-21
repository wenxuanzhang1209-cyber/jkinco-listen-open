from pathlib import Path
import zipfile

from docx import Document

from report_templates import build_report_docx


ROOT = Path(__file__).resolve().parents[1]
TALK_TEMPLATE = ROOT / "templates" / "v2" / "会议纪要_工地例会原版式模板.docx"
LOGO = ROOT / "assets" / "sribs-meeting-logo.jpeg"


def document_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_engineering_template_is_clean_and_keeps_only_logo():
    text = document_text(TALK_TEMPLATE)
    assert "{{PROJECT_NAME}}" in text
    assert "{{REPORT_BODY}}" in text
    for sample_value in ("上中路 466 号", "第十六次", "2024 年"):
        assert sample_value not in text
    with zipfile.ZipFile(TALK_TEMPLATE) as archive:
        images = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert len(images) == 1


def test_all_scene_docx_exports(tmp_path):
    summaries = {
        "talk": "# 会议纪要\n项目名称：示例建设项目\n会议名称：质量安全例会\n会议内容：完成验收资料归档。",
        "general": "# 会议纪要\n## 重点事项\n风险闭环完成。",
        "personal": "# 个人备忘录\n## 待办事项\n|事项|状态|\n|---|---|\n|提交周报|进行中|",
        "interview": "# 面试记录与候选人反馈表\n## 能力评价\n候选人岗位匹配度较高。",
        "customer_visit": "# 客户拜访会议纪要\n## 客户诉求\n客户希望开展试点。",
        "auto": "# 智能识别场景路由\n目标场景：会议纪要",
    }
    for mode, summary in summaries.items():
        output = tmp_path / f"{mode}.docx"
        build_report_docx(summary, mode, output, talk_template=TALK_TEMPLATE, logo=LOGO)
        assert output.exists() and output.stat().st_size > 10_000
