"""自定义模板导出时,纪要正文只能出现一次。

模板里带 {{minutes}} 且插入策略选「追加到新页」时,原实现会把原始 markdown
填进占位符,同时又在文末追加一份格式化的 —— 导出的 Word 里纪要出现两遍,
其中一遍带着 ## 和 - 符号。三种插入策略里只有 append 会中,平时用 auto 测不出来。
"""
import io
import sys
import tempfile
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from backend.custom_templates import analyze_docx, render_custom_docx, validate_docx

DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUMMARY = "## 议题\n- 第一条\n- 第二条"


def _template(paragraphs):
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render(content, strategy, target, fields=None):
    analysis = analyze_docx(content, validate_docx(content, "t.docx", DOCX_TYPE))
    if not target:
        target = analysis.get("recommended_target", "append:new-page")
    with tempfile.TemporaryDirectory() as directory:
        output = Path(directory) / "out.docx"
        render_custom_docx(content, SUMMARY, output, fields=fields or {"title": "周例会"},
                           analysis=analysis, insertion_strategy=strategy, insertion_target=target)
        return [p.text for p in Document(str(output)).paragraphs if p.text.strip()]


@pytest.mark.parametrize("strategy", ["auto", "append", "manual"])
def test_minutes_appear_exactly_once(strategy):
    content = _template(["会议主题：{{title}}", "{{minutes}}"])
    analysis = analyze_docx(content, validate_docx(content, "t.docx", DOCX_TYPE))
    target = "append:new-page" if strategy == "append" else (
        (analysis.get("insertion_candidates") or [{}])[0].get("id", "")
    )
    lines = _render(content, strategy, target)
    assert lines.count("议题") == 1, f"{strategy}: 纪要出现 {lines.count('议题')} 次 -> {lines}"
    assert lines.count("第一条") == 1


@pytest.mark.parametrize("strategy", ["auto", "append", "manual"])
def test_raw_markdown_never_reaches_the_document(strategy):
    """导出的 Word 里不该出现 ## 和 - 这类 markdown 符号。"""
    content = _template(["会议主题：{{title}}", "{{minutes}}"])
    analysis = analyze_docx(content, validate_docx(content, "t.docx", DOCX_TYPE))
    target = "append:new-page" if strategy == "append" else (
        (analysis.get("insertion_candidates") or [{}])[0].get("id", "")
    )
    lines = _render(content, strategy, target)
    raw = [line for line in lines if line.lstrip().startswith(("##", "- "))]
    assert not raw, f"{strategy}: 输出里残留原始 markdown {raw}"


def test_unfilled_fields_still_show_the_placeholder_hint():
    """纪要占位符置空不能连累其它字段:未填写的字段仍要显示「待确认」。"""
    content = _template(["地点：{{location}}", "{{minutes}}"])
    lines = _render(content, "append", "append:new-page", fields={"title": "周例会"})
    assert any("地点：待确认" in line for line in lines), lines
