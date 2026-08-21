"""访客的配额必须单独收紧。

访客免注册,单 IP 在锁定窗口内能开 5 个,而账号保留到会话有效期的两倍(8 小时)
——保留期内单 IP 最多可存活两千多个访客账号。模板配额却是按账号给的(20 个 /
40MB),没有任何按访客收紧的一档,放大出来的总量足以写满磁盘,而磁盘写满时整个
平台一起停摆,不只是模板功能。

同一代码库对作业槽位早就区分了访客(MAX_PENDING_JOBS_PER_GUEST),理由写得很
清楚:「按账号计的配额对访客形同虚设」。模板这条路漏了同一道。

另一件事:GUEST_MAX_PER_WINDOW 此前定义了却从未被使用,访客接口走的是
login_blocked 的默认上限 LOGIN_MAX_FAILURES。两者当前恰好都是 5,所以没有行为
差异 —— 但那是个陷阱:为宽容密码手误调大 JKINCO_LOGIN_MAX_FAILURES 会同时放宽
访客开号,而调 JKINCO_GUEST_MAX_PER_WINDOW 毫无效果。
"""
from __future__ import annotations

import inspect

import backend.auth as auth
import backend.custom_templates as templates
import backend.main as main


def test_guest_template_quota_is_tighter():
    assert templates.MAX_TEMPLATES_PER_GUEST < templates.MAX_TEMPLATES_PER_USER
    assert templates.MAX_TEMPLATE_STORAGE_PER_GUEST < templates.MAX_TEMPLATE_STORAGE_PER_USER


def test_the_quota_is_actually_applied_by_role():
    """常量存在不等于生效 —— 这一轮已经撞到过一次(GUEST_MAX_PER_WINDOW)。"""
    source = inspect.getsource(templates.create_template)
    assert "_template_quota_for(owner)" in source, "上传时没有按身份取配额"
    quota = inspect.getsource(templates._template_quota_for)
    assert "is_guest" in quota and "MAX_TEMPLATES_PER_GUEST" in quota


def test_guest_gets_the_guest_numbers(monkeypatch):
    monkeypatch.setattr(templates, "is_guest", lambda name: True, raising=False)
    import backend.auth as auth_module
    monkeypatch.setattr(auth_module, "is_guest", lambda name: True)
    assert templates._template_quota_for("guest_abc") == (
        templates.MAX_TEMPLATES_PER_GUEST, templates.MAX_TEMPLATE_STORAGE_PER_GUEST,
    )


def test_registered_user_keeps_the_full_quota(monkeypatch):
    import backend.auth as auth_module
    monkeypatch.setattr(auth_module, "is_guest", lambda name: False)
    assert templates._template_quota_for("alice") == (
        templates.MAX_TEMPLATES_PER_USER, templates.MAX_TEMPLATE_STORAGE_PER_USER,
    )


# --- 那个从未生效的常量 ---

def test_guest_window_limit_is_wired_up():
    source = inspect.getsource(main.guest_access)
    assert "login_blocked(guest_key, GUEST_MAX_PER_WINDOW)" in source, (
        "访客开号仍走 login_blocked 的默认上限,GUEST_MAX_PER_WINDOW 是死代码"
    )


def test_the_two_knobs_are_no_longer_coupled(monkeypatch):
    """调大登录失败容忍度,不该顺带放宽访客开号。"""
    monkeypatch.setattr(auth, "LOGIN_MAX_FAILURES", 99)
    key = "test-guest-decoupled"
    auth.LOGIN_FAILURES.pop(key, None)
    for _ in range(auth.GUEST_MAX_PER_WINDOW):
        auth.record_login_failure(key)
    assert auth.login_blocked(key, auth.GUEST_MAX_PER_WINDOW) is True, "访客上限没有独立生效"
    assert auth.login_blocked(key) is False, "前提校验:默认上限此时确实更宽松"
    auth.LOGIN_FAILURES.pop(key, None)


# --- 死锁看门狗 ---
# 按身份取配额需要查 platform_users,而那要 PROFILE_DB_LOCK;create_template 自己
# 也持有同一把锁。第一版把取配额写在锁内 —— PROFILE_DB_LOCK 是普通 Lock 不是
# RLock,于是直接自锁死,整个测试套件停在这条上,栈里是 _pthread_cond_wait。
#
# 挂起是最差的失败方式:不报错、不结束,CI 只会一直等。所以这条用例带看门狗,
# 让同类问题以「失败」而不是「卡住」的形式暴露出来。

def test_create_template_does_not_deadlock():
    import io
    import threading

    from docx import Document

    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("模板")
    document.add_paragraph("{{minutes}}")
    document.save(buffer)

    done = threading.Event()
    reached_quota = threading.Event()

    # 只有真的走到配额检查,这条用例才有意义 —— 第一版把参数个数写错,
    # TypeError 被 except 吞掉、done 立刻置位,用例永远通过。
    original_quota = templates._template_quota_for

    def traced(owner):
        reached_quota.set()
        return original_quota(owner)

    def run():
        try:
            templates.create_template("deadlock_probe", "探针模板", "probe.docx", buffer.getvalue(),
                                      scenario="general")
        except BaseException:                     # 配额/校验类报错不算死锁
            pass
        finally:
            done.set()

    templates._template_quota_for = traced
    try:
        worker = threading.Thread(target=run, daemon=True)
        worker.start()
        assert done.wait(timeout=20), (
            "create_template 20 秒未返回 —— 极可能又在持锁状态下调用了需要同一把锁的函数"
        )
    finally:
        templates._template_quota_for = original_quota
    assert reached_quota.is_set(), "没走到配额检查,这条用例是空转的"


def test_quota_lookup_happens_before_the_lock():
    """结构上钉住:取配额必须在 with PROFILE_DB_LOCK 之前。"""
    source = inspect.getsource(templates.create_template)
    quota_at = source.index("_template_quota_for(owner)")
    lock_at = source.index("with PROFILE_DB_LOCK")
    assert quota_at < lock_at, "取配额又被挪进锁里了,会自锁死"
