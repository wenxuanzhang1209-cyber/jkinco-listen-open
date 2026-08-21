"""自定义纪要模板的安全边界与导出回归测试。"""
from __future__ import annotations

import base64
from concurrent.futures import ThreadPoolExecutor
import hashlib
import io
from pathlib import Path
import zipfile

import pytest
from docx import Document
from fastapi.testclient import TestClient
from PIL import Image

from backend.main import app
from backend.custom_templates import (
    analyze_docx,
    generate_minutes_with_template,
    render_custom_docx,
    validate_docx,
)
from helpers import solve_captcha


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _walk_table_text(table) -> list[str]:
    values: list[str] = []
    seen: set[object] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_element = cell._tc
            if cell_element in seen:
                continue
            seen.add(cell_element)
            values.extend(paragraph.text for paragraph in cell.paragraphs)
            for nested in cell.tables:
                values.extend(_walk_table_text(nested))
    return values


def _all_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        values.extend(_walk_table_text(table))
    seen: set[int] = set()
    for section in document.sections:
        for story in (section.header, section.first_page_header, section.even_page_header,
                      section.footer, section.first_page_footer, section.even_page_footer):
            identity = id(story._element)
            if identity in seen:
                continue
            seen.add(identity)
            values.extend(paragraph.text for paragraph in story.paragraphs)
            for table in story.tables:
                values.extend(_walk_table_text(table))
    return "\n".join(values)


def _replace_zip_entry(content: bytes, name: str, replacement: bytes) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(content)) as source, zipfile.ZipFile(
        output, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            target.writestr(item, replacement if item.filename == name else source.read(item.filename))
    return output.getvalue()


def _complex_template_bytes() -> bytes:
    document = Document()
    title = document.add_paragraph()
    first = title.add_run("{{会")
    first.bold = True
    title.add_run("议主题}}")
    document.add_paragraph("以下固定说明必须保留。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "固定项目资料"
    table.cell(1, 0).text = "参会人员"
    table.cell(1, 1).text = "{{参会人员}}"
    image_bytes = io.BytesIO()
    Image.new("RGB", (32, 18), "#1374e8").save(image_bytes, format="PNG")
    image_bytes.seek(0)
    document.add_picture(image_bytes)
    document.add_paragraph("正文开始 {{会议纪要}} 固定落款")
    document.add_paragraph("{{待办事项}}")
    header = document.sections[0].header.paragraphs[0]
    header.add_run("会议日期：{{会")
    header.add_run("议日期}}")
    document.sections[0].footer.paragraphs[0].text = "生成于 {{生成时间}}"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _register(client: TestClient, username: str) -> None:
    challenge = client.get("/api/auth/captcha").json()
    answer = solve_captcha(challenge["token"])
    response = client.post("/api/auth/register", json={
        "username": username,
        "display_name": username,
        "password": "StrongPass123",
        "captcha_token": challenge["token"],
        "captcha_answer": answer,
    })
    assert response.status_code == 201, response.text


def _upload(client: TestClient, filename: str = "项目周会模板.docx") -> dict:
    response = client.post(
        "/api/custom-templates",
        data={"name": "项目周会模板"},
        files={"file": (filename, _docx_bytes("项目周会", "{{会议主题}}", "{{会议纪要}}"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 201, response.text
    return response.json()


def test_template_crud_is_private_between_accounts():
    with TestClient(app) as owner:
        _register(owner, "template_owner")
        created = _upload(owner)
        assert [item["id"] for item in owner.get("/api/custom-templates").json()["items"]] == [created["id"]]

    with TestClient(app) as other:
        _register(other, "template_other")
        assert other.get("/api/custom-templates").json()["items"] == []
        assert other.delete(f"/api/custom-templates/{created['id']}").status_code == 404
        assert other.post(
            "/api/process",
            data={
                "live_text": "讨论本周项目计划。",
                "process_mode": "只转写，不推送",
                "app_mode": "general",
                "custom_template_id": created["id"],
            },
        ).status_code == 404
        assert other.post(
            "/api/export/docx",
            json={"summary": "私有纪要", "mode": "general", "custom_template_id": created["id"]},
        ).status_code == 404

    with TestClient(app) as owner:
        assert owner.post("/api/auth/login", json={"username": "template_owner", "password": "StrongPass123"}).status_code == 200
        assert owner.delete(f"/api/custom-templates/{created['id']}").status_code == 204
        assert owner.get("/api/custom-templates").json()["items"] == []


def test_upload_rejects_non_docx_and_corrupt_archive():
    with TestClient(app) as client:
        _register(client, "template_invalid")
        wrong_type = client.post(
            "/api/custom-templates",
            files={"file": ("template.txt", b"not a docx", "text/plain")},
        )
        assert wrong_type.status_code == 400
        corrupt = client.post(
            "/api/custom-templates",
            files={"file": ("template.docx", b"not a zip", "application/octet-stream")},
        )
        assert corrupt.status_code == 400


def test_validation_rejects_macro_and_path_traversal():
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", b"application/vnd.ms-word.document.macroEnabled.main+xml")
        archive.writestr("word/document.xml", b"<document/>")
        archive.writestr("../escape", b"x")
    try:
        validate_docx(buffer.getvalue(), "unsafe.docx")
        raise AssertionError("unsafe archive was accepted")
    except ValueError as error:
        assert "不安全路径" in str(error) or "宏" in str(error)


def test_custom_export_replaces_marker_and_keeps_template_content():
    with TestClient(app) as client:
        _register(client, "template_export")
        created = _upload(client)
        response = client.post(
            "/api/export/docx",
            json={
                "summary": "# 交付周会\n确认周五完成联调。",
                "mode": "general",
                "custom_template_id": created["id"],
            },
        )
        assert response.status_code == 200, response.text
        exported = Document(io.BytesIO(response.content))
        text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        assert "项目周会" in text
        assert "交付周会" in text
        assert "确认周五完成联调" in text
        assert "{{会议纪要}}" not in text


def test_markerless_template_appends_minutes(tmp_path):
    target = tmp_path / "rendered.docx"
    render_custom_docx(_docx_bytes("固定封面标题"), "# 会议纪要\n形成三项行动计划。", target)
    rendered = Document(target)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "固定封面标题" in text
    assert "形成三项行动计划" in text


def test_template_generation_reformats_existing_minutes(monkeypatch):
    captured = {}

    def fake_llm(prompt: str, **_kwargs):
        captured["prompt"] = prompt
        return "# 项目周会\n- 已完成"

    monkeypatch.setattr("backend.custom_templates.call_llm", fake_llm)
    template = {"name": "周会模板", "content": _docx_bytes("一、上周完成", "二、本周计划")}
    result = generate_minutes_with_template("# 原纪要\n确认已完成联调。", template)
    assert result.startswith("# 项目周会")
    assert "已核验会议纪要" in captured["prompt"]
    assert "确认已完成联调" in captured["prompt"]
    assert "原始转写" not in captured["prompt"]


def test_split_run_placeholders_across_body_table_header_footer_and_image(tmp_path):
    source = _complex_template_bytes()
    analysis = analyze_docx(source)
    locations = {(item["field"], item["part"], item["kind"]) for item in analysis["placeholders"]}
    assert ("title", "body", "paragraph") in locations
    assert ("participants", "body", "table_cell") in locations
    assert any(field == "date" and part.startswith("header:") for field, part, _kind in locations)
    assert any(field == "generated_at" and part.startswith("footer:") for field, part, _kind in locations)
    assert analysis["recommended_target"].startswith("placeholder:")

    target = tmp_path / "complex-rendered.docx"
    render_custom_docx(
        source,
        "# 交付协调会\n确认联调窗口，固定落款必须保留。",
        target,
        fields={
            "title": "交付协调会",
            "date": "2026-07-30",
            "participants": ["张三", "李四"],
            "todos": [
                {
                    "task": "完成接口联调",
                    "owner": "张三",
                    "deadline": "2026-08-01",
                    "status": "进行中",
                },
                {
                    "task": "归档验收资料",
                    "owner": "李四",
                    "deadline": "2026-08-02",
                    "status": "未开始",
                },
            ],
            "generated_at": "2026-07-30 12:00:00",
        },
        analysis=analysis,
    )

    rendered = Document(target)
    text = _all_text(rendered)
    assert "以下固定说明必须保留" in text
    assert "固定项目资料" in text
    assert "交付协调会" in text
    assert "2026-07-30" in text
    assert "张三\n李四" in text
    assert "确认联调窗口" in text
    assert "固定落款" in text
    assert "待办事项" in text
    assert "负责人" in text
    assert "截止时间" in text
    assert "完成接口联调" in text
    assert "生成于 2026-07-30 12:00:00" in text
    assert "{{" not in text
    assert "${" not in text
    assert "undefined" not in text.lower()
    assert "null" not in text.lower()
    assert rendered.paragraphs[0].runs[0].bold is True
    with zipfile.ZipFile(target) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())


def test_semantic_auto_insertion_precedes_fixed_tail(tmp_path):
    source = _docx_bytes("固定封面", "会议主要内容", "固定签名：上海建科咨询")
    analysis = analyze_docx(source)
    assert analysis["recommended_target"] == "after:body/p1"
    assert analysis["parse_status"] == "ready"

    target = tmp_path / "semantic.docx"
    render_custom_docx(
        source,
        "# 会议纪要\n- 已明确交付节点\n- 已确认责任人",
        target,
        analysis=analysis,
    )
    texts = [paragraph.text for paragraph in Document(target).paragraphs if paragraph.text.strip()]
    assert texts.index("会议主要内容") < texts.index("会议纪要")
    assert texts.index("已确认责任人") < texts.index("固定签名：上海建科咨询")


def test_manual_insertion_allows_plain_structure_node(tmp_path):
    source = _docx_bytes("封面说明", "固定落款")
    analysis = analyze_docx(source)
    candidate_ids = {item["id"] for item in analysis["insertion_candidates"]}
    assert "after:body/p0" in candidate_ids
    assert analysis["recommended_target"] == "append:new-page"

    target = tmp_path / "manual.docx"
    render_custom_docx(
        source,
        "# 人工确认纪要\n放在封面说明之后。",
        target,
        analysis=analysis,
        insertion_strategy="manual",
        insertion_target="after:body/p0",
    )
    texts = [paragraph.text for paragraph in Document(target).paragraphs if paragraph.text.strip()]
    assert texts.index("封面说明") < texts.index("人工确认纪要")
    assert texts.index("放在封面说明之后。") < texts.index("固定落款")


def test_markerless_low_confidence_requires_confirmation_and_safe_append(tmp_path):
    source = _docx_bytes("公司固定说明", "负责人签字：________")
    analysis = analyze_docx(source)
    assert analysis["parse_status"] == "needs_confirmation"
    assert analysis["recommended_target"] == "append:new-page"
    assert any("确认插入位置" in message for message in analysis["risk_messages"])

    target = tmp_path / "safe-append.docx"
    render_custom_docx(source, "# 通用会议纪要\n原模板内容不得覆盖。", target, analysis=analysis)
    texts = [paragraph.text for paragraph in Document(target).paragraphs if paragraph.text.strip()]
    assert texts[:2] == ["公司固定说明", "负责人签字：________"]
    assert texts.count("通用会议纪要") == 1
    assert texts.count("原模板内容不得覆盖。") == 1


def test_repeated_generation_starts_from_pristine_source_and_is_concurrent(tmp_path):
    source = _docx_bytes("模板固定内容", "{{会议纪要}}", "固定落款")
    source_digest = source
    analysis = analyze_docx(source)

    def render(index: int) -> str:
        target = tmp_path / f"concurrent-{index}.docx"
        render_custom_docx(
            source,
            "# 并发纪要\n本次生成只能出现一次。",
            target,
            analysis=analysis,
        )
        return _all_text(Document(target))

    with ThreadPoolExecutor(max_workers=4) as executor:
        results = list(executor.map(render, range(8)))

    assert source == source_digest
    assert len(set(results)) == 1
    for text in results:
        assert text.count("并发纪要") == 1
        assert text.count("本次生成只能出现一次。") == 1
        assert "模板固定内容" in text
        assert "固定落款" in text


def test_validation_rejects_mime_compression_bomb_dangerous_xml_and_external_parts():
    source = _docx_bytes("安全模板")
    try:
        validate_docx(source, "safe.docx", "text/plain")
        raise AssertionError("mismatched MIME type was accepted")
    except ValueError as error:
        assert "类型" in str(error)

    bomb = io.BytesIO()
    with zipfile.ZipFile(bomb, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"<Types/>")
        archive.writestr("word/document.xml", b"A" * (2 * 1024 * 1024))
    try:
        validate_docx(bomb.getvalue(), "bomb.docx")
        raise AssertionError("compression bomb was accepted")
    except ValueError as error:
        assert "压缩比例" in str(error)

    unsafe_xml = _replace_zip_entry(
        source,
        "word/document.xml",
        b'<!DOCTYPE x [<!ENTITY secret SYSTEM "file:///etc/passwd">]><x>&secret;</x>',
    )
    try:
        validate_docx(unsafe_xml, "xxe.docx")
        raise AssertionError("unsafe XML was accepted")
    except ValueError as error:
        assert "实体" in str(error)

    relationship_name = "word/_rels/document.xml.rels"
    with zipfile.ZipFile(io.BytesIO(source)) as archive:
        relationships = archive.read(relationship_name)
    external = relationships.replace(
        b"</Relationships>",
        (
            b'<Relationship Id="rExternal" '
            b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
            b'Target="http://127.0.0.1:1/tracker.png" TargetMode="External"/>'
            b"</Relationships>"
        ),
    )
    unsafe_relationship = _replace_zip_entry(source, relationship_name, external)
    try:
        validate_docx(unsafe_relationship, "external.docx")
        raise AssertionError("external relationship was accepted")
    except ValueError as error:
        assert "外部资源" in str(error)


def test_template_settings_default_download_soft_delete_and_historical_export():
    with TestClient(app) as client:
        _register(client, "template_management")
        first = _upload(client, "first.docx")
        second_response = client.post(
            "/api/custom-templates",
            data={"name": "第二模板", "scenario": "general"},
            files={
                "file": (
                    "second.docx",
                    _docx_bytes("第二模板", "{{会议纪要}}"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert second_response.status_code == 201, second_response.text
        second = second_response.json()

        updated_first = client.patch(
            f"/api/custom-templates/{first['id']}",
            json={
                "name": "第一默认模板",
                "scenario": "general",
                "is_default": True,
                "insertion_strategy": "auto",
                "insertion_target": first["analysis"]["recommended_target"],
            },
        )
        assert updated_first.status_code == 200, updated_first.text
        updated_second = client.patch(
            f"/api/custom-templates/{second['id']}",
            json={
                "scenario": "general",
                "is_default": True,
                "insertion_strategy": "auto",
                "insertion_target": second["analysis"]["recommended_target"],
            },
        )
        assert updated_second.status_code == 200, updated_second.text
        items = client.get("/api/custom-templates").json()["items"]
        defaults = [item for item in items if item["scenario"] == "general" and item["is_default"]]
        assert [item["id"] for item in defaults] == [second["id"]]

        downloaded = client.get(f"/api/custom-templates/{first['id']}/download")
        assert downloaded.status_code == 200
        assert hashlib.sha256(downloaded.content).hexdigest() == first["sha256"]

        assert client.delete(f"/api/custom-templates/{first['id']}").status_code == 204
        assert client.get(f"/api/custom-templates/{first['id']}").status_code == 404
        historical_export = client.post(
            "/api/export/docx",
            json={
                "summary": "# 历史纪要\n删除模板后，既有记录仍可复现。",
                "mode": "general",
                "custom_template_id": first["id"],
            },
        )
        assert historical_export.status_code == 200, historical_export.text
        text = _all_text(Document(io.BytesIO(historical_export.content)))
        assert "删除模板后，既有记录仍可复现" in text


# --- 自定义模板导出:控制字符与属性长度 ---
# 「控制字符让 Word 导出永久失败」修过一次,但只修在 clean_markdown_text 覆盖的
# summary/minutes 上。自定义模板是另一条导出路径,字段值走 _field_text ——
# title、participants、location、transcript,以及模型生成的 overview(映射到
# {{summary}} 占位符)全都绕开了清洗,同一个 ValueError 原样复现。
#
# overview 由模型产出并落库:里面一旦混进一个控制字符,这份纪要用自定义模板
# 就再也导不出来 —— 和当初那次的「永久失败」是同一回事。

CONTROL_CHAR_FIELDS = [
    ("overview", "会议概览\x0b要点", "模型生成的概览,映射到 {{summary}}"),
    ("participants", "张三\x00李四", "参会人员"),
    ("transcript", "发言\x1f内容", "原始转写"),
    ("title", "项目\x0b评审", "会议主题"),
    ("location", "三楼\x08会议室", "会议地点"),
]


@pytest.mark.parametrize("field, value, label", CONTROL_CHAR_FIELDS)
def test_control_characters_in_any_field_do_not_break_export(field, value, label, tmp_path):
    template = _docx_bytes("概览：{{summary}}", "参会：{{participants}}", "地点：{{location}}",
                           "转写：{{transcript}}", "主题：{{title}}", "{{minutes}}")
    output = tmp_path / f"{field}.docx"
    render_custom_docx(template, "## 结论\n- 通过", output, fields={field: value})
    Document(str(output))  # 必须能被重新打开


def test_emoji_and_rare_characters_survive_the_cleaning(tmp_path):
    """清洗不能顺手把 BMP 之外的字符删掉 —— 那是另一次修复的教训。"""
    template = _docx_bytes("参会：{{participants}}", "{{minutes}}")
    output = tmp_path / "emoji.docx"
    render_custom_docx(template, "## 结论\n- 通过 🎉", output, fields={"participants": "张三 🎉 𠮷"})
    body = "\n".join(p.text for p in Document(str(output)).paragraphs)
    assert "🎉" in body and "𠮷" in body, "emoji/生僻字被误删了"


def test_a_long_first_line_does_not_break_export(tmp_path):
    """标题为空时会从纪要首行推导,而 docx 核心属性有 255 字符硬上限。

    首行不受任何长度约束 —— 模型输出一段长开头,整份导出就抛
    ValueError("exceeded 255 char limit for property")。
    """
    template = _docx_bytes("模板", "{{minutes}}")
    output = tmp_path / "long.docx"
    long_first_line = "本次会议围绕" + "各专业分项的进度与质量控制措施" * 20 + "展开讨论"
    assert len(long_first_line) > 255
    render_custom_docx(template, f"{long_first_line}\n\n## 结论\n- 通过", output)
    document = Document(str(output))
    assert len(document.core_properties.title) <= 255
    body = "\n".join(p.text for p in document.paragraphs)
    assert long_first_line in body, "正文里的完整首行不该被属性截断连累"
